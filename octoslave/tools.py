import atexit
import base64
import json
import os
import posixpath
import re
import shutil
import signal
import subprocess
import tempfile
import threading
import time as _time
import glob as glob_module
from pathlib import Path

from .tools_bio import BIO_TOOL_DEFINITIONS, BIO_TOOL_NAMES, execute_bio_tool
from .tools_cryo import (
    CRYO_TOOL_DEFINITIONS,
    CRYO_TOOL_NAMES,
    CRYO_MODIFYING_TOOLS,
    execute_cryo_tool,
)

# Optional web deps — imported lazily inside functions
try:
    import requests as _requests
    _HAS_REQUESTS = True
except ImportError:
    _HAS_REQUESTS = False

try:
    from ddgs import DDGS as _DDGS
    _HAS_DDG = True
except ImportError:
    try:
        from duckduckgo_search import DDGS as _DDGS
        _HAS_DDG = True
    except ImportError:
        _HAS_DDG = False

try:
    from bs4 import BeautifulSoup as _BS4
    _HAS_BS4 = True
except ImportError:
    _HAS_BS4 = False

try:
    from playwright.sync_api import sync_playwright as _sync_playwright
    _HAS_PLAYWRIGHT = True
except ImportError:
    _HAS_PLAYWRIGHT = False

# Tools that require permission in controlled mode
MODIFYING_TOOLS = {"write_file", "edit_file", "apply_patch", "bash",
                   "run_background", "stop_process"} | set(CRYO_MODIFYING_TOOLS)

# Tools that require permission only in controlled mode (not supervised)
FILE_MODIFYING_TOOLS = {"write_file", "edit_file", "apply_patch"}

# ---------------------------------------------------------------------------
# Tool schemas (sent to the model)
# ---------------------------------------------------------------------------

TOOL_DEFINITIONS = [
    {
        "type": "function",
        "function": {
            "name": "read_file",
            "description": (
                "Read the contents of a file. Always read a file before editing it. "
                "Returns line-numbered content."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Absolute or relative file path"},
                    "offset": {"type": "integer", "description": "Start line (1-indexed, optional)"},
                    "limit": {"type": "integer", "description": "Max lines to read (optional)"},
                },
                "required": ["path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "write_file",
            "description": (
                "Create a new file or completely overwrite an existing file. "
                "Prefer edit_file for modifying existing files."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "File path to write"},
                    "content": {"type": "string", "description": "Full file content"},
                },
                "required": ["path", "content"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "edit_file",
            "description": (
                "Make a targeted edit by replacing an exact string with a new string. "
                "By default old_string must appear exactly once (uniquely). "
                "Set replace_all=true to replace every occurrence — useful for renames "
                "and global refactors. Preserve indentation and surrounding context exactly."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "File path to edit"},
                    "old_string": {"type": "string", "description": "Exact string to replace. Must be unique unless replace_all is true."},
                    "new_string": {"type": "string", "description": "Replacement string"},
                    "replace_all": {"type": "boolean", "description": "Replace every occurrence instead of requiring uniqueness (default false)"},
                },
                "required": ["path", "old_string", "new_string"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "bash",
            "description": (
                "Execute a shell command and return stdout + stderr. "
                "Use for running tests, installing packages, building, git operations, etc. "
                "Do NOT start long-running servers or blocking processes (e.g. 'python app.py', 'flask run') — "
                "they will block until timeout. To verify a web app, use syntax checks or import tests instead. "
                "Default timeout is 300 s. "
                "For any job that may exceed a few minutes (model training, large simulations/resampling "
                "runs, big data processing, dev servers), PREFER run_background over raising this timeout — a long "
                "bash call blocks everything (in a team, the whole lab stalls) until it returns. Only raise "
                "timeout for bounded one-offs like a large pip install or a full test suite (e.g. 600)."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "command": {"type": "string", "description": "Shell command to run"},
                    "timeout": {"type": "integer", "description": "Timeout in seconds (default 300). Prefer run_background over a large timeout for long/expensive jobs; only raise this for bounded one-offs like big installs or full test suites (600-3600)."},
                },
                "required": ["command"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "glob",
            "description": "Find files matching a glob pattern. Returns list of matching paths.",
            "parameters": {
                "type": "object",
                "properties": {
                    "pattern": {"type": "string", "description": "Glob pattern, e.g. '**/*.py' or 'src/**/*.ts'"},
                    "path": {"type": "string", "description": "Root directory (defaults to working dir)"},
                },
                "required": ["pattern"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "grep",
            "description": "Search file contents with a regex pattern. Returns matching file:line:content entries.",
            "parameters": {
                "type": "object",
                "properties": {
                    "pattern": {"type": "string", "description": "Regex pattern to search"},
                    "path": {"type": "string", "description": "File or directory to search (defaults to working dir)"},
                    "glob": {"type": "string", "description": "File filter, e.g. '*.py'"},
                    "case_insensitive": {"type": "boolean", "description": "Case-insensitive search"},
                },
                "required": ["pattern"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "list_dir",
            "description": "List the contents of a directory with file sizes.",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Directory path (defaults to working dir)"},
                },
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "web_search",
            "description": (
                "Search the web using DuckDuckGo. Returns titles, URLs, and snippets. "
                "Use for finding research papers, documentation, current events, or any web information. "
                "Then use web_fetch to read the full content of promising URLs."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {"type": "string", "description": "Search query"},
                    "max_results": {"type": "integer", "description": "Max results to return (default 10, max 20)"},
                    "region": {"type": "string", "description": "Region code e.g. 'us-en', 'wt-wt' (worldwide). Default 'wt-wt'"},
                },
                "required": ["query"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "web_fetch",
            "description": (
                "Fetch and extract readable text content from a URL. "
                "Strips HTML tags, navigation, and ads to return the main text. "
                "Use after web_search to read full articles, papers, or documentation pages."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "url": {"type": "string", "description": "URL to fetch"},
                    "max_chars": {"type": "integer", "description": "Max characters to return (default 8000)"},
                },
                "required": ["url"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "crawl_tree",
            "description": (
                "Crawl a website tree starting from a root URL using BFS. "
                "Follows internal links up to max_depth levels deep and returns a JSON tree "
                "of {url: {title, text_snippet, children, depth}}. "
                "Uses Playwright (handles JS-rendered pages) if available, otherwise requests+BeautifulSoup. "
                "Ideal for scraping category trees, product hierarchies, documentation trees, or any "
                "multi-level site structure. Save results with output_path."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "root_url": {"type": "string", "description": "Starting URL to crawl from"},
                    "link_selector": {"type": "string", "description": "CSS selector for links to follow (default 'a')"},
                    "url_pattern": {"type": "string", "description": "Regex — only follow URLs matching this pattern (optional)"},
                    "max_depth": {"type": "integer", "description": "Maximum depth to crawl (default 3)"},
                    "max_pages": {"type": "integer", "description": "Maximum number of pages to visit (default 100)"},
                    "same_domain": {"type": "boolean", "description": "Only follow links on the same domain (default true)"},
                    "use_js": {"type": "boolean", "description": "Force Playwright even for static sites (default false — auto-selects best engine)"},
                    "output_path": {"type": "string", "description": "If given, save full JSON tree to this file path"},
                },
                "required": ["root_url"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "compress_log",
            "description": (
                "Compress a long log, training output, or noisy command output into a compact "
                "templated summary using the `codag` CLI. Drastically reduces context cost "
                "(typically 95–99% token reduction) while preserving rare events like errors "
                "and tracebacks verbatim. "
                "USE THIS instead of read_file or bash whenever you would otherwise read a log "
                "longer than ~500 lines: ML training logs, kubectl/docker logs, CI output, "
                "stack traces from a failed run, anything tee'd to disk. "
                "Provide exactly ONE of `command` (codag runs and wraps it) or `path` "
                "(codag reads a local file). Default mode `compact` is plain text, free, "
                "no auth. Mode `capsule` returns structured JSON but requires `codag auth login`."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "command": {"type": "string", "description": "Shell command whose output should be wrapped and compressed (e.g. 'python train.py --epochs 10'). Mutually exclusive with `path`."},
                    "path": {"type": "string", "description": "Path to an existing log file to compress (absolute or relative to working dir). Mutually exclusive with `command`."},
                    "service": {"type": "string", "description": "Optional tag for these lines (e.g. 'train', 'api'). Helps codag's template inference."},
                    "mode": {"type": "string", "description": "Output mode: 'compact' (default, plain text, no auth) or 'capsule' (JSON, requires auth)."},
                    "timeout": {"type": "integer", "description": "Max seconds for the wrapped command. Default 600 for compress_log of a file, matches the inner command's needs otherwise."},
                },
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "image_ocr",
            "description": (
                "Extract text from an image file (PNG, JPG/JPEG, TIFF, BMP, GIF, WEBP) "
                "using tesseract OCR. Use for screenshots, scanned documents, photos of "
                "text, figures with embedded labels, plots with axis ticks, etc. "
                "For PDFs use pdf_ocr instead. "
                "Requires pytesseract + the `tesseract` binary "
                "(macOS: `brew install tesseract`, Debian: `apt install tesseract-ocr`)."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Path to image file"},
                    "lang": {"type": "string", "description": "Tesseract language code (default 'eng'). Use 'eng+deu' for multilingual."},
                    "preprocess": {"type": "boolean", "description": "Apply grayscale + threshold for noisy/low-contrast images (default false)"},
                    "psm": {"type": "integer", "description": "Tesseract page segmentation mode (default 3 = auto). Try 6 for a single block of text, 11 for sparse text."},
                },
                "required": ["path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "view_image",
            "description": (
                "LOOK at an image yourself — the actual pixels are attached to the "
                "conversation so you can see them. Use this for anything where the "
                "picture is the information: plots and charts (read trends, outliers, "
                "where curves cross, whether a fit is bad), molecular structures, gels "
                "and blots, micrographs, UI screenshots you must judge visually, "
                "rendered figures you just produced. "
                "This is NOT image_ocr: image_ocr extracts text with tesseract and is "
                "the right tool for scanned documents, screenshots of text, and reading "
                "exact printed numbers/labels. Use both on the same figure when you "
                "need the shape AND the precise tick values. "
                "The image arrives as a separate message right after this tool's reply. "
                "Requires a vision-capable model — if the active model has no image "
                "input the call fails and says so; it never silently drops the image."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Path to the image file (PNG/JPG/TIFF/BMP/GIF/WEBP), absolute or relative to the working dir. Large images are downscaled automatically."},
                    "note": {"type": "string", "description": "Optional: what you are looking for (e.g. 'does the red curve plateau?'). Attached alongside the image to keep your own intent in view."},
                },
                "required": ["path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "apply_patch",
            "description": (
                "Apply MULTIPLE find-and-replace edits to a SINGLE file in one call. "
                "Each edit replaces an exact old_string with a new_string; edits apply "
                "in order and the file is written only if ALL edits succeed (atomic). "
                "Much cheaper than several edit_file calls when changing several regions "
                "of the same file. Each old_string must be unique in the file when its "
                "edit runs, unless that edit sets replace_all=true. Preserve indentation "
                "and surrounding context exactly. Pass `edits` as a JSON ARRAY of objects "
                '(e.g. [{"old_string": "...", "new_string": "..."}]) — not a string and '
                "not a single object."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "File to patch"},
                    "edits": {
                        "type": "array",
                        "description": "Edits applied in order",
                        "items": {
                            "type": "object",
                            "properties": {
                                "old_string": {"type": "string", "description": "Exact text to replace"},
                                "new_string": {"type": "string", "description": "Replacement text"},
                                "replace_all": {"type": "boolean", "description": "Replace every occurrence for this edit (default false)"},
                            },
                            "required": ["old_string", "new_string"],
                        },
                    },
                },
                "required": ["path", "edits"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "todo_write",
            "description": (
                "Create or update a task checklist for the current work. Pass the FULL "
                "list every time — it REPLACES the previous list. Use it on any multi-step "
                "task: keep exactly one item 'in_progress', mark items 'completed' as you "
                "finish them, and add new items as they emerge. This keeps you and the user "
                "oriented on long tasks and shows live progress in the UI. Skip it for "
                "trivial single-step requests."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "todos": {
                        "type": "array",
                        "description": "The complete, ordered task list",
                        "items": {
                            "type": "object",
                            "properties": {
                                "content": {"type": "string", "description": "Imperative description of the step"},
                                "status": {"type": "string", "enum": ["pending", "in_progress", "completed"]},
                            },
                            "required": ["content", "status"],
                        },
                    },
                },
                "required": ["todos"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "run_background",
            "description": (
                "Start a long-running or blocking command in the BACKGROUND and return "
                "immediately with a process id. Use this — not bash — for dev servers, "
                "training jobs, file watchers, or anything that does not exit on its own "
                "(the bash tool blocks until the command finishes). Output is captured to a "
                "file; read it with check_process and stop the job with stop_process."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "command": {"type": "string", "description": "Shell command to run in the background"},
                    "cwd": {"type": "string", "description": "Working directory (defaults to the session working dir)"},
                },
                "required": ["command"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "check_process",
            "description": (
                "Check a background process started with run_background: returns its "
                "running/exited status, exit code, and recent output. Call with no id to "
                "list all background processes."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "id": {"type": "string", "description": "Process id from run_background (omit to list all)"},
                    "tail_lines": {"type": "integer", "description": "Trailing output lines to return (default 50)"},
                },
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "stop_process",
            "description": "Stop a background process started with run_background (SIGTERM, then SIGKILL).",
            "parameters": {
                "type": "object",
                "properties": {
                    "id": {"type": "string", "description": "Process id from run_background"},
                },
                "required": ["id"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "ask_user",
            "description": (
                "Ask the human a clarifying question and wait for their answer. Use ONLY "
                "when genuinely blocked on a decision that is the user's to make and cannot "
                "be resolved from the task, the code, or sensible defaults — not for routine "
                "choices you can make yourself. In autonomous or non-interactive runs the "
                "user may be unavailable; if no answer comes back, proceed with your best "
                "judgement."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "question": {"type": "string", "description": "The question to ask"},
                    "options": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": "Optional suggested answers the user can pick from",
                    },
                },
                "required": ["question"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "remember",
            "description": (
                "Save a durable insight to this PROJECT's memory — a fact, lesson, "
                "preference, or gotcha that will help in FUTURE sessions in this "
                "same project, not just this one. Use it sparingly and "
                "deliberately when you learn "
                "something worth carrying forward: a non-obvious project quirk, a "
                "fix that took effort to find, a stable user preference, an "
                "environment constraint (e.g. a tool that isn't installed), or a "
                "dead end worth avoiding next time. Write the insight so it makes "
                "sense out of context, with specifics. Do NOT use it for routine "
                "progress, transient state, or things already obvious from the code "
                "or the task. These notes are recalled automatically when relevant. "
                "Re-stating an existing fact UPDATES it (the old version is replaced) "
                "— keep memory concise."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "content": {"type": "string", "description": "The insight to remember, self-contained and specific."},
                    "tags": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": "Optional short tags for retrieval (e.g. ['build', 'macos']).",
                    },
                },
                "required": ["content"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "forget",
            "description": (
                "Remove an insight from this project's memory that is no longer true "
                "or relevant (e.g. a fact that changed, a quirk that was fixed, advice "
                "that proved wrong). Describe the insight to drop; the closest "
                "matching remembered insight(s) are removed. Use this to keep memory "
                "accurate and lean — don't let stale notes accumulate."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "content": {"type": "string", "description": "A description of the insight to remove (matched against stored insights)."},
                },
                "required": ["content"],
            },
        },
    },
] + BIO_TOOL_DEFINITIONS


# ---------------------------------------------------------------------------
# MCP (Model Context Protocol) — user-wired external tools
# ---------------------------------------------------------------------------

def init_mcp(cfg: dict | None = None, force: bool = False,
             working_dir: str | None = None) -> None:
    """Connect to configured MCP servers (idempotent). Safe to call from every
    entry point — single-agent REPL, research pipeline, web UI.

    ``working_dir`` seeds the MCP roots so filesystem-style servers start out
    sandboxed to the task's directory (kept in sync per-call in ``execute_tool``).
    """
    try:
        from .mcp_client import manager
        roots = [working_dir] if working_dir else None
        manager.init_from_config(cfg, force=force, roots=roots)
    except Exception:
        # MCP is strictly additive — never let it break the core toolbox.
        pass


def _mcp_manager():
    try:
        from .mcp_client import manager
        return manager
    except Exception:
        return None


# Curated tool surface for small local models. Handing a 12B model all ~37
# schemas (incl. bio lookups + MCP log-tailers) dilutes its attention and
# induces wrong/malformed tool calls. Local runs see only the essentials for
# file/data/web work. Big models (e-INFRA / NIM / custom) are unaffected — they
# only ever pass profile=None and keep the full surface.
LOCAL_TOOL_ALLOWLIST = frozenset({
    "read_file", "write_file", "edit_file", "bash",
    "glob", "grep", "list_dir",
    "web_search", "web_fetch",
    # Local vision models (llava, gemma3, qwen-vl, minicpm-v) are a common
    # reason to run Ollama at all. Safe on a text-only local model too: Ollama's
    # /api/show reports the vision capability, so view_image refuses outright
    # rather than handing over an image the runner would silently drop.
    "view_image",
    # ...and the fallback that refusal points at MUST be reachable here. Most
    # local models are text-only, so this is exactly the profile where OCR is
    # the only way to get at an image at all — granting view_image without it
    # would leave a blind model with no route to the picture.
    "image_ocr", "pdf_ocr",
})


# ---------------------------------------------------------------------------
# Dynamic tool registry — runtime-built tools (the Lab "foundry") and the
# Lab meta-tools register here. Each entry maps a tool name to (definition,
# func) where func has signature func(args: dict, working_dir: str) ->
# tuple[str, bool]. Strictly additive: empty by default, never affects the
# core toolbox or non-Lab callers.
# ---------------------------------------------------------------------------
_DYNAMIC_TOOLS: dict[str, "tuple[dict, callable]"] = {}


def register_dynamic_tool(definition: dict, func) -> None:
    """Register (or replace) a runtime tool. ``definition`` is an OpenAI
    function-tool schema; ``func(args, working_dir) -> (text, ok)``."""
    name = definition["function"]["name"]
    _DYNAMIC_TOOLS[name] = (definition, func)


def unregister_dynamic_tool(name: str) -> None:
    _DYNAMIC_TOOLS.pop(name, None)


def dynamic_tool_names() -> set[str]:
    return set(_DYNAMIC_TOOLS)


def all_tool_definitions(profile: str | None = None) -> list[dict]:
    """Built-in tools plus any tools exposed by connected MCP servers, plus any
    runtime-registered dynamic tools (Lab foundry / meta-tools).

    When profile == "local" (small local/Ollama models), return only the
    curated LOCAL_TOOL_ALLOWLIST and skip MCP tools entirely — a tight,
    high-signal surface that small models can actually use reliably.
    """
    if profile == "local":
        return [td for td in TOOL_DEFINITIONS
                if td["function"]["name"] in LOCAL_TOOL_ALLOWLIST]
    defs = list(TOOL_DEFINITIONS)
    # CryoUncle: the CryoSPARC toolbox is exposed ONLY for the cryouncle profile
    # so other profiles aren't handed a dozen extra cryo-EM schemas. When called
    # without an explicit profile, fall back to this thread's active profile.
    if _cryo_enabled(profile if profile is not None else current_tool_profile()):
        defs += CRYO_TOOL_DEFINITIONS
    mgr = _mcp_manager()
    if mgr is not None:
        defs += mgr.tool_definitions()
    defs += [d for (d, _f) in _DYNAMIC_TOOLS.values()]
    return defs


def valid_tool_names() -> set[str]:
    """Names of all callable tools (built-in + MCP + dynamic). Used by the
    text-format tool-call fallback parser."""
    names = {td["function"]["name"] for td in TOOL_DEFINITIONS}
    # Cryo tool names only enter the (fallback-parser) namespace under cryouncle,
    # so no other profile is contaminated by them.
    if _cryo_enabled():
        names |= set(CRYO_TOOL_NAMES)
    mgr = _mcp_manager()
    if mgr is not None:
        names |= mgr.tool_names()
    names |= set(_DYNAMIC_TOOLS)
    return names


# ---------------------------------------------------------------------------
# Execution backend (local by default, or a remote host over SSH)
#
# A per-thread flag mirrors agent._RT: when a remote target is configured for the
# current agent thread, the filesystem/bash helpers route their work over an SSH
# session instead of the local machine. The default (no remote) path is byte-for-
# byte unchanged, so Local mode carries zero behavior change.
# ---------------------------------------------------------------------------

_EXEC = threading.local()

# Bio/OCR tools that read a `path` input and/or write output files. In remote
# mode these run against a local temp mirror (input staged in, new outputs
# pushed back) — they need local Python libs (rdkit/pymupdf/pandas) that we do
# not assume exist on the remote host.
_STAGED_TOOLS = frozenset(BIO_TOOL_NAMES) | {"image_ocr", "view_image"}


def configure_execution(remote_cfg: dict | None) -> None:
    """Point this thread's tools at a remote host (dict) or back to local (None)."""
    _EXEC.remote_cfg = remote_cfg


def set_tool_profile(profile: str | None) -> None:
    """Record the active prompt profile for THIS agent thread. Profile-scoped
    toolboxes (currently the CryoSPARC / cryouncle tools) key off this so they
    are both invisible AND inert for every other profile. Set by
    agent.configure_runtime at the start of each run; defaults to None (no
    profile-scoped tools) for any entry point that never sets it."""
    _EXEC.tool_profile = profile


def current_tool_profile() -> str | None:
    return getattr(_EXEC, "tool_profile", None)


def _cryo_enabled(profile: str | None = None) -> bool:
    """CryoSPARC tools are available only under the cryouncle profile."""
    if profile is None:
        profile = current_tool_profile()
    return profile == "cryouncle"


def _remote():
    """Return the RemoteSession for this thread, or None for local execution."""
    # While a staged remote file is being processed with the LOCAL helpers
    # (read_file / bio / OCR on a pulled temp copy), routing is forced off so
    # those helpers don't re-route back to the remote and recurse.
    if getattr(_EXEC, "force_local", False):
        return None
    cfg = getattr(_EXEC, "remote_cfg", None)
    if not cfg:
        return None
    try:
        from .remote import RemoteSession
        return RemoteSession.get(cfg)
    except Exception:
        return None


def remote_label() -> str | None:
    """Human label for the active remote target, or None when local."""
    cfg = getattr(_EXEC, "remote_cfg", None)
    if not cfg:
        return None
    return cfg.get("name") or cfg.get("host") or cfg.get("id")


def active_remote_cfg() -> dict | None:
    """The remote config this thread's tools route to, or None when local. Lets a
    prompt builder describe the active remote without threading it through every
    call (execution routing is already thread-local via configure_execution)."""
    return getattr(_EXEC, "remote_cfg", None)


def _remote_resolve(path: str, working_dir: str) -> str:
    """Resolve ``path`` against the remote working dir using POSIX semantics.

    Never uses local ``Path.resolve`` (which would mangle a remote path on a
    machine where it does not exist).
    """
    if not path:
        return posixpath.normpath(working_dir)
    if posixpath.isabs(path):
        return posixpath.normpath(path)
    return posixpath.normpath(posixpath.join(working_dir, path))


# local working_dir → real remote directory, per remote host (probe/mkdir once).
_REMOTE_WD_MAP: dict[tuple, str] = {}


def _remote_cwd(working_dir: str) -> str:
    """Map a session's (local) working_dir to a REAL directory on the active
    remote host.

    In whole-session-remote mode the agent's working_dir is a LOCAL absolute path
    that does not exist on the remote, so ``cd``/``mkdir`` of it fail (or, worse,
    try to create ``/private`` etc.). Instead we run in
    ``<remote_dir or $HOME>/octoslave/<basename>`` — created once and cached. A
    working_dir that already exists on the remote (the user gave a real remote
    path) is used unchanged. No-op when execution is local."""
    sess = _remote()
    if sess is None or not working_dir:
        return working_dir
    cfg = active_remote_cfg() or {}
    key = (cfg.get("id") or cfg.get("host") or "", working_dir)
    hit = _REMOTE_WD_MAP.get(key)
    if hit:
        return hit
    try:
        if posixpath.isabs(working_dir) and sess.is_dir(working_dir):
            _REMOTE_WD_MAP[key] = working_dir
            return working_dir
        base = (cfg.get("remote_dir") or "").strip() or sess.home()
        name = posixpath.basename(working_dir.rstrip("/")) or "session"
        rwd = posixpath.normpath(posixpath.join(base, "octoslave", name))
        sess.mkdirs(rwd)
        _REMOTE_WD_MAP[key] = rwd
        return rwd
    except Exception:
        return working_dir


def _format_bash_output(stdout: str, stderr: str, returncode: int) -> str:
    """Shared stdout/stderr/exit-code formatting for local + remote bash."""
    stdout = stdout or ""
    stderr = stderr or ""
    if stdout and stderr:
        output = f"--- stdout ---\n{stdout}\n--- stderr ---\n{stderr}"
    else:
        output = stdout or stderr
    if not output:
        output = f"(exit code {returncode})"
    elif returncode != 0:
        output += f"\n(exit code {returncode})"
    if len(output) > 8000:
        output = output[:2000] + "\n\n... [output truncated] ...\n\n" + output[-5000:]
    return output


# ---- remote tool implementations ------------------------------------------

def _remote_read_file(sess, path, working_dir, offset=None, limit=None) -> tuple[str, bool]:
    rp = _remote_resolve(path, working_dir)
    if not sess.exists(rp):
        return f"File not found: {path}", False
    if not sess.is_file(rp):
        return f"Not a file: {path}", False
    suffix = posixpath.splitext(rp)[1].lower()
    size = sess.size(rp)
    # Large plain-text file with no window → preview head remotely (no full pull).
    if (offset is None and limit is None and size > _LARGE_TEXT_BYTES
            and suffix not in (".pdf", ".xlsx", ".xls", ".xlsm", ".ods", ".docx", ".doc")):
        import shlex
        head_out, _, _ = sess.run(f"head -n {_LARGE_TEXT_PREVIEW_LINES} {shlex.quote(rp)}", None, timeout=60)
        preview_lines = head_out.splitlines()
        numbered = "\n".join(f"{i + 1}\t{ln}" for i, ln in enumerate(preview_lines))
        if suffix in (".csv", ".tsv", ".tab", ".parquet", ".pq", ".jsonl", ".ndjson"):
            tip = f"This is a data table — call bio_inspect(path='{path}') for a schema-aware preview."
        else:
            tip = "Use bash (head/grep/awk) or read_file with offset/limit to page through it; do not read it whole."
        return (
            f"File: {path} ({size:,} bytes — too large to read whole; showing first "
            f"{len(preview_lines)} lines).\n{numbered}\n\n[LARGE FILE] {tip}",
            True,
        )
    # Stage the file locally and reuse the full local reader (PDF/Excel/docx/binary).
    try:
        local = sess.pull(rp, suffix=suffix)
    except Exception as e:
        return f"Could not read remote file {path}: {e}", False
    _EXEC.force_local = True
    try:
        res, ok = _read_file(local, os.path.dirname(local), offset, limit)
    finally:
        _EXEC.force_local = False
        try:
            os.unlink(local)
        except OSError:
            pass
    # The local reader labels output with the temp path — restore the real path.
    res = res.replace(f"File: {local}", f"File: {path}", 1)
    res = res.replace(local, path)
    return res, ok


def _remote_write_file(sess, path, content, working_dir) -> tuple[str, bool]:
    rp = _remote_resolve(path, working_dir)
    size = len(content.encode("utf-8"))
    if size > _MAX_DOWNLOAD_BYTES:
        return (f"Refused to write {path} — content is {size / (1024 ** 3):.2f} GB, "
                f"exceeds the 1 GB limit.", False)
    try:
        sess.write_text(rp, content)
    except Exception as e:
        return f"Could not write remote file {path}: {e}", False
    lines = content.count("\n") + 1
    return f"Written {lines} lines to {path}", True


def _remote_edit_file(sess, path, old_string, new_string, working_dir, replace_all=False) -> tuple[str, bool]:
    rp = _remote_resolve(path, working_dir)
    if not sess.is_file(rp):
        return f"File not found: {path}", False
    if old_string == "":
        return "old_string must not be empty. To create a file, use write_file.", False
    if old_string == new_string:
        return "old_string and new_string are identical — no change needed.", False
    try:
        content = sess.read_text(rp)
    except Exception as e:
        return f"Could not read remote file {path}: {e}", False
    count = content.count(old_string)
    if count == 0:
        return f"String not found in {path}:\n{old_string[:200]}", False
    if count > 1 and not replace_all:
        return (f"old_string appears {count} times in {path} — make it more specific to "
                f"ensure uniqueness, or pass replace_all=true to replace every occurrence.", False)
    new_content = content.replace(old_string, new_string)
    try:
        sess.write_text(rp, new_content)
    except Exception as e:
        return f"Could not write remote file {path}: {e}", False
    if replace_all and count > 1:
        return f"Edited {path} ({count} occurrences replaced)", True
    return f"Edited {path}", True


def _remote_apply_patch(sess, path, edits, working_dir) -> tuple[str, bool]:
    rp = _remote_resolve(path, working_dir)
    if not sess.is_file(rp):
        return f"File not found: {path}", False
    edits = _coerce_edits(edits)  # idempotent when already a clean list
    if not edits:
        return (
            "apply_patch got no usable edits. Pass `edits` as a JSON array of "
            '{old_string, new_string} objects — e.g. edits=[{"old_string": "foo", '
            '"new_string": "bar"}]. (file left UNCHANGED)'
        ), False
    try:
        content = sess.read_text(rp)
    except Exception as e:
        return f"Could not read remote file {path}: {e}", False
    applied = 0
    notes: list[str] = []
    for i, e in enumerate(edits, 1):
        if not isinstance(e, dict):
            return f"Edit #{i} is not an object — expected {{old_string, new_string}}.", False
        old = e.get("old_string", "")
        new = e.get("new_string", "")
        replace_all = bool(e.get("replace_all", False))
        if old == "":
            return f"Edit #{i}: old_string must not be empty. (no changes written)", False
        if old == new:
            return f"Edit #{i}: old_string and new_string are identical. (no changes written)", False
        count = content.count(old)
        if count == 0:
            return (f"Edit #{i}: string not found:\n{old[:200]}\n"
                    f"(applied {applied}/{len(edits)} so far — file left UNCHANGED)", False)
        if count > 1 and not replace_all:
            return (f"Edit #{i}: old_string appears {count} times — make it unique or set "
                    f"replace_all=true for this edit. (file left UNCHANGED)", False)
        content = content.replace(old, new)
        applied += 1
        if replace_all and count > 1:
            notes.append(f"#{i}×{count}")
    try:
        sess.write_text(rp, content)
    except Exception as e:
        return f"Could not write remote file {path}: {e}", False
    suffix = f" ({', '.join(notes)} replace_all)" if notes else ""
    return f"Applied {applied} edit(s) to {path}{suffix}", True


def _remote_bash(sess, command, working_dir, timeout=300) -> tuple[str, bool]:
    stdout, stderr, rc = sess.run(command, cwd=working_dir, timeout=timeout)
    return _format_bash_output(stdout, stderr, rc), rc == 0


def _remote_glob(sess, pattern, working_dir, path=None) -> tuple[str, bool]:
    import shlex
    if posixpath.isabs(pattern):
        root = posixpath.dirname(pattern) or "/"
        name = posixpath.basename(pattern)
    else:
        root = _remote_resolve(path, working_dir) if path else working_dir
        name = pattern
    recursive = name.startswith("**/")
    leaf = name[3:] if recursive else name
    if "/" in leaf or recursive:
        cmd = f"find {shlex.quote(root)} -name {shlex.quote(posixpath.basename(leaf))} -print 2>/dev/null | head -200"
    else:
        cmd = f"find {shlex.quote(root)} -maxdepth 1 -name {shlex.quote(leaf)} -print 2>/dev/null | head -200"
    out, _, _ = sess.run(cmd, None, timeout=60)
    matches = [m for m in out.splitlines() if m.strip()]
    if not matches:
        return f"No files matching '{pattern}'", True
    result = []
    for m in matches:
        result.append(m[len(working_dir):].lstrip("/") if m.startswith(working_dir) else m)
    output = "\n".join(result)
    if len(matches) >= 200:
        output += "\n... (showing first 200)"
    return output, True


def _remote_grep(sess, pattern, working_dir, path=None, glob=None, case_insensitive=False) -> tuple[str, bool]:
    import shlex
    flags = "-rn"
    if case_insensitive:
        flags += "i"
    inc = f"--include={shlex.quote(glob)} " if glob else ""
    target = _remote_resolve(path, working_dir) if path else working_dir
    cmd = f"grep {flags} {inc}-e {shlex.quote(pattern)} {shlex.quote(target)}"
    out, err, _ = sess.run(cmd, None, timeout=60)
    output = out or err or "No matches found"
    if len(output) > 8000:
        lines = output.splitlines()
        output = "\n".join(lines[:150]) + f"\n... ({len(lines)} total matches, showing first 150)"
    return output, True


def _remote_list_dir(sess, working_dir, path=None) -> tuple[str, bool]:
    import shlex
    target = _remote_resolve(path, working_dir) if path else working_dir
    if not sess.exists(target):
        return f"Directory not found: {path}", False
    if not sess.is_dir(target):
        return f"Not a directory: {path}", False
    out, _, _ = sess.run(f"ls -1Ap {shlex.quote(target)}", None, timeout=60)
    entries = [e for e in out.splitlines() if e.strip()]
    dirs = sorted(e for e in entries if e.endswith("/"))
    files = sorted(e for e in entries if not e.endswith("/"))
    lines = [f"  {d}" for d in dirs] + [f"  {f}" for f in files]
    return f"{target}/\n" + "\n".join(lines), True


def _run_staged_tool(name, args, sess, working_dir) -> tuple[str, bool]:
    """Run a bio/OCR tool against a local temp mirror of the remote workdir.

    Input file (``path``) is pulled in; any new files the tool writes are pushed
    back to the remote working dir. Lets tools that need local Python libs work
    even though the data lives remotely.
    """
    tmp = tempfile.mkdtemp(prefix="ots_stage_")
    try:
        local_args = dict(args)
        val = args.get("path")
        if isinstance(val, str) and val:
            rp = _remote_resolve(val, working_dir)
            if sess.is_file(rp):
                base = posixpath.basename(rp)
                try:
                    sess.pull(rp, dest=os.path.join(tmp, base))
                    local_args["path"] = base
                except Exception as e:
                    return f"Could not stage remote file {val}: {e}", False
        before = set(os.listdir(tmp))
        _EXEC.force_local = True
        try:
            if name == "image_ocr":
                result, ok = _image_ocr(working_dir=tmp, **local_args)
            elif name == "view_image":
                # Label with the path the model asked for, not the temp basename
                # we staged it to. `label` is ours to set, so a model that also
                # passed one must not collide with it.
                local_args.pop("label", None)
                result, ok = _view_image(working_dir=tmp, label=str(val or ""), **local_args)
            else:
                result, ok = execute_bio_tool(name, local_args, tmp)
        finally:
            _EXEC.force_local = False
        # Push back any new top-level files the tool produced.
        for fn in os.listdir(tmp):
            if fn in before:
                continue
            full = os.path.join(tmp, fn)
            if os.path.isfile(full):
                try:
                    sess.push(full, _remote_resolve(fn, working_dir))
                except Exception:
                    pass
        return result, ok
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


# ---------------------------------------------------------------------------
# Tool execution
# ---------------------------------------------------------------------------

def execute_tool(name: str, args: dict, working_dir: str, permission_mode: str = "autonomous") -> tuple[str, bool]:
    """Execute a tool. Returns (result_text, success)."""
    # A stop already requested must not start NEW work. Without this, a turn that
    # proposed several tool calls would run all of them after the user hit Stop —
    # the single biggest reason stopping felt slow.
    from . import interrupt
    if interrupt.should_stop():
        raise interrupt.StopRequested
    # Check permission for modifying tools
    # - controlled: ask for all modifying tools (file ops + bash)
    # - supervised: ask only for file operations (allow bash without asking)
    _is_mcp = name.startswith("mcp__")
    # MCP tools call out to external servers (which may write files, hit APIs,
    # mutate state) — gate them like modifying tools in controlled mode.
    if permission_mode == "controlled" and (name in MODIFYING_TOOLS or _is_mcp):
        try:
            from . import display
            if not display.request_permission(name, args, working_dir, permission_mode):
                return f"Permission denied by user for {name}", False
        except Exception:
            # Fallback if display module has issues
            pass
    elif permission_mode == "supervised" and name in FILE_MODIFYING_TOOLS:
        try:
            from . import display
            if not display.request_permission(name, args, working_dir, permission_mode):
                return f"Permission denied by user for {name}", False
        except Exception:
            # Fallback if display module has issues
            pass
    
    # Reject empty / missing required string args BEFORE dispatch — some
    # models (notably nemotron-49b acting as Orchestrator) emit tool calls
    # with `path=""` etc. and a generic "File not found:" reply does not
    # help them recover. Return an actionable message instead.
    err = _validate_required_args(name, args)
    if err is not None:
        return err, False

    # When the whole session runs on a remote host, translate the local session
    # working_dir into a real remote directory so every tool (bash, file ops,
    # cluster-job submission) resolves paths and `cd`s somewhere that exists.
    # No-op for local execution.
    orig_working_dir = working_dir  # pre-remap, for a safe retry (see except below)
    if _remote() is not None:
        working_dir = _remote_cwd(working_dir)

    try:
        if name == "read_file":
            return _read_file(working_dir=working_dir, **args)
        elif name == "write_file":
            return _write_file(working_dir=working_dir, **args)
        elif name == "edit_file":
            return _edit_file(working_dir=working_dir, **args)
        elif name == "apply_patch":
            return _apply_patch(working_dir=working_dir, **args)
        elif name == "todo_write":
            return _todo_write(working_dir=working_dir, **args)
        elif name == "bash":
            return _bash(working_dir=working_dir, **args)
        elif name == "run_background":
            return _run_background(working_dir=working_dir, **args)
        elif name == "check_process":
            return _check_process(working_dir=working_dir, **args)
        elif name == "stop_process":
            return _stop_process(working_dir=working_dir, **args)
        elif name == "ask_user":
            return _ask_user(working_dir=working_dir, **args)
        elif name == "remember":
            return _remember(working_dir=working_dir, **args)
        elif name == "forget":
            return _forget(working_dir=working_dir, **args)
        elif name == "glob":
            return _glob(working_dir=working_dir, **args)
        elif name == "grep":
            return _grep(working_dir=working_dir, **args)
        elif name == "list_dir":
            return _list_dir(working_dir=working_dir, **args)
        elif name == "web_search":
            return _web_search(**args)
        elif name == "web_fetch":
            return _web_fetch(**args)
        elif name == "crawl_tree":
            return _crawl_tree(**args)
        elif name == "compress_log":
            return _compress_log(working_dir=working_dir, **args)
        elif name == "image_ocr":
            _sess = _remote()
            if _sess is not None:
                return _run_staged_tool("image_ocr", args, _sess, working_dir)
            return _image_ocr(working_dir=working_dir, **args)
        elif name == "view_image":
            _sess = _remote()
            if _sess is not None:
                return _run_staged_tool("view_image", args, _sess, working_dir)
            return _view_image(working_dir=working_dir, **args)
        elif name in BIO_TOOL_NAMES:
            _sess = _remote()
            if _sess is not None:
                return _run_staged_tool(name, args, _sess, working_dir)
            return execute_bio_tool(name, args, working_dir)
        elif name in CRYO_TOOL_NAMES:
            # Hard isolation: cryo tools only run under the cryouncle profile, so
            # no other profile can invoke them even via a hallucinated call.
            if not _cryo_enabled():
                return (f"{name} is only available under the 'cryouncle' prompt "
                        f"profile (CryoSPARC companion). Switch with /profile "
                        f"cryouncle or -p cryouncle.", False)
            # CryoSPARC tools talk to the user's instance over the network, so
            # they run locally (no remote staging) and reuse the cached client.
            return execute_cryo_tool(name, args, working_dir)
        elif _is_mcp:
            mgr = _mcp_manager()
            if mgr is None:
                return f"MCP is unavailable; cannot run {name}", False
            # Keep filesystem-style MCP servers sandboxed to the *current* task
            # workdir (users pick workdirs per task/message). Idempotent + cheap
            # when unchanged; pushes roots/list_changed when it actually moves.
            # Skip when executing remotely — the local MCP server can't see the
            # remote working dir, so pinning a root to a remote path is meaningless.
            if working_dir and _remote() is None:
                try:
                    mgr.set_roots([working_dir])
                except Exception:
                    pass
            return mgr.call(name, args)
        elif name in _DYNAMIC_TOOLS:
            _def, func = _DYNAMIC_TOOLS[name]
            return func(args, working_dir)
        else:
            return f"Unknown tool: {name}", False
    except TypeError as e:
        # Some models hallucinate extra keyword args they've seen in OTHER agent
        # frameworks (e.g. bash with description=/mode=, borrowed from Claude
        # Code). Rather than hard-failing the call, strip the single offending
        # kwarg and retry — the action still runs. Bounded: each retry removes one
        # key that is actually present in args, so it converges. Re-enters with the
        # ORIGINAL (pre-remap) working_dir so the remote path isn't remapped twice.
        m = re.search(r"unexpected keyword argument '([^']+)'", str(e))
        if m and m.group(1) in args:
            cleaned = {k: v for k, v in args.items() if k != m.group(1)}
            return execute_tool(name, cleaned, orig_working_dir, permission_mode)
        return f"Invalid arguments for {name}: {e}", False
    except Exception as e:
        return f"Tool error: {e}", False


# Required string args per tool — used by _validate_required_args.
# Tools accepting `working_dir` as the only "input" (list_dir) intentionally
# allow path="" to mean "the working dir itself".
_REQUIRED_STR_ARGS: dict[str, tuple[str, ...]] = {
    "read_file": ("path",),
    "write_file": ("path", "content"),
    "edit_file": ("path", "old_string", "new_string"),
    "apply_patch": ("path",),
    "bash": ("command",),
    "run_background": ("command",),
    "stop_process": ("id",),
    "ask_user": ("question",),
    "remember": ("content",),
    "forget": ("content",),
    "glob": ("pattern",),
    "grep": ("pattern",),
    "web_search": ("query",),
    "web_fetch": ("url",),
    "crawl_tree": ("root_url",),
    "image_ocr": ("path",),
    "view_image": ("path",),
    # bio tools that need a string input
    "bio_inspect": ("path",),
    "rdkit_describe": ("smiles",),
    "pdb_fetch": ("pdb_id",),
    "alphafold_fetch": ("uniprot_id",),
    "ena_fetch": ("accession",),
    "pdf_ocr": ("path",),
    # CryoUncle / CryoSPARC tools
    "cryo_workspaces": ("project_uid",),
    "cryo_jobs": ("project_uid",),
    "cryo_job": ("project_uid", "job_uid"),
    "cryo_dataset": ("project_uid", "job_uid"),
    "cryo_create_job": ("project_uid", "workspace_uid", "type"),
    "cryo_queue_job": ("project_uid", "job_uid"),
    "cryo_control_job": ("project_uid", "job_uid", "action"),
    "cryo_download": ("project_uid", "path"),
}


def _validate_required_args(name: str, args: dict) -> str | None:
    """Return an actionable error string if a required string arg is missing/empty, else None."""
    if not isinstance(args, dict):
        return f"Tool `{name}` requires JSON object arguments, got {type(args).__name__}."
    required = _REQUIRED_STR_ARGS.get(name)
    if not required:
        return None
    missing: list[str] = []
    empty: list[str] = []
    for key in required:
        # `content` for write_file may legitimately be empty (creating empty file)
        if name == "write_file" and key == "content":
            if key not in args:
                missing.append(key)
            continue
        if key not in args:
            missing.append(key)
            continue
        v = args[key]
        if v is None or (isinstance(v, str) and v.strip() == ""):
            empty.append(key)
    if not missing and not empty:
        return None
    parts = []
    if missing:
        parts.append(f"missing argument(s): {', '.join(missing)}")
    if empty:
        parts.append(f"empty argument(s): {', '.join(empty)}")
    hint = _ARG_HINTS.get(name, "")
    return (
        f"Cannot run `{name}` — {'; '.join(parts)}. "
        f"{hint} "
        f"Re-issue the call with the required argument filled in."
    ).strip()


_ARG_HINTS: dict[str, str] = {
    "read_file": "Pass `path` as an absolute or working-dir-relative file path "
                 "(e.g. {\"path\": \"round_001/01_literature.md\"}).",
    "write_file": "Pass `path` and `content` (e.g. {\"path\": \"round_001/06_synthesis.md\", \"content\": \"...\"}).",
    "edit_file": "Pass `path`, `old_string`, `new_string`. Use replace_all=true for renames.",
    "apply_patch": "Pass `path` and `edits` (a list of {old_string, new_string} objects).",
    "bash": "Pass `command` as a shell string (e.g. {\"command\": \"ls\"}).",
    "run_background": "Pass `command` as a shell string (e.g. {\"command\": \"python app.py\"}).",
    "stop_process": "Pass `id` as the process id returned by run_background (e.g. \"bg1\").",
    "ask_user": "Pass `question` as a clear, specific question string.",
    "remember": "Pass `content` as a self-contained insight string (e.g. {\"content\": \"This repo uses uv, not pip.\"}).",
    "forget": "Pass `content` describing the stale insight to remove (e.g. {\"content\": \"repo uses uv not pip\"}).",
    "glob": "Pass `pattern` (e.g. {\"pattern\": \"**/*.py\"}).",
    "grep": "Pass `pattern` as a regex (e.g. {\"pattern\": \"def \\\\w+\"}).",
    "web_search": "Pass `query`.",
    "web_fetch": "Pass `url` as a fully-qualified http(s) URL.",
    "crawl_tree": "Pass `root_url` as a fully-qualified http(s) URL.",
    "image_ocr": "Pass `path` to an image file (PNG/JPG/TIFF/BMP/GIF/WEBP).",
    "view_image": "Pass `path` to an image file (PNG/JPG/TIFF/BMP/GIF/WEBP) to look at it yourself.",
    "pdf_ocr": "Pass `path` to a PDF file. Use `pages` to limit page range.",
    "bio_inspect": "Pass `path` to a data table (CSV/TSV/Parquet/JSONL) or bio/chem data file.",
    "rdkit_describe": "Pass `smiles` (e.g. {\"smiles\": \"CC(=O)O\"}).",
    "pdb_fetch": "Pass `pdb_id` (4-char RCSB ID, e.g. \"1CRN\").",
    "alphafold_fetch": "Pass `uniprot_id` (e.g. \"P12345\").",
    "ena_fetch": "Pass `accession` (e.g. \"SRR000001\").",
}


def _resolve(path: str, working_dir: str) -> Path:
    p = Path(path)
    if not p.is_absolute():
        p = Path(working_dir) / p
    return p.resolve()


def _is_binary(path: Path) -> bool:
    """Quick check: read first 512 bytes and look for null bytes or high non-printable ratio."""
    try:
        chunk = path.read_bytes()[:512]
        if b"\x00" in chunk:
            return True
        non_printable = sum(1 for b in chunk if b < 9 or (14 <= b < 32))
        return non_printable / max(len(chunk), 1) > 0.30
    except OSError:
        return False


def _extract_pdf(resolved: Path, offset: int = None, limit: int = None) -> tuple[str, bool]:
    try:
        import fitz  # pymupdf
        doc = fitz.open(str(resolved))
        parts = [f"--- Page {i+1} ---\n{page.get_text()}" for i, page in enumerate(doc)]
        num_pages = len(doc)
    except ImportError:
        return "pymupdf not installed. Run: pip install pymupdf", False
    except Exception as e:
        return f"Could not open PDF: {e}", False

    if not any(p.strip() for p in parts):
        return "PDF contains no extractable text (may be scanned/image-based).", False

    full_text = "\n\n".join(parts)
    lines = full_text.splitlines()
    total = len(lines)

    start = (offset - 1) if offset and offset > 0 else 0
    end = (start + limit) if limit else total
    selected = lines[start:end]

    header = f"PDF: {resolved.name} ({num_pages} pages, {total} lines extracted)\n\n"
    return header + "\n".join(selected), True


def _extract_excel(resolved: Path, offset: int = None, limit: int = None) -> tuple[str, bool]:
    try:
        import openpyxl
    except ImportError:
        try:
            import subprocess, sys
            subprocess.check_call([sys.executable, "-m", "pip", "install", "openpyxl", "-q"])
            import openpyxl
        except Exception:
            return "openpyxl not installed. Run: pip install openpyxl", False

    try:
        wb = openpyxl.load_workbook(str(resolved), read_only=True, data_only=True)
    except Exception as e:
        return f"Could not open Excel file: {e}", False

    parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            # Skip entirely empty rows
            if any(cell is not None for cell in row):
                rows.append("\t".join("" if v is None else str(v) for v in row))
        if rows:
            parts.append(f"=== Sheet: {sheet_name} ===\n" + "\n".join(rows))

    if not parts:
        return f"Excel file {resolved.name} appears empty.", False

    full_text = "\n\n".join(parts)
    lines = full_text.splitlines()
    total = len(lines)

    start = (offset - 1) if offset and offset > 0 else 0
    end = (start + limit) if limit else total
    selected = lines[start:end]

    header = f"Excel: {resolved.name} ({len(wb.sheetnames)} sheet(s), {total} lines)\n\n"
    return header + "\n".join(selected), True


def _extract_docx(resolved: Path, offset: int = None, limit: int = None) -> tuple[str, bool]:
    try:
        import docx as _docx
    except ImportError:
        try:
            import subprocess, sys
            subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
            import docx as _docx
        except Exception:
            return "python-docx not installed. Run: pip install python-docx", False

    try:
        doc = _docx.Document(str(resolved))
    except Exception as e:
        return f"Could not open Word document: {e}", False

    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)

    if not parts:
        return f"Word document {resolved.name} appears empty.", False

    full_text = "\n".join(parts)
    lines = full_text.splitlines()
    total = len(lines)

    start = (offset - 1) if offset and offset > 0 else 0
    end = (start + limit) if limit else total
    selected = lines[start:end]

    header = f"Word: {resolved.name} ({total} lines extracted)\n\n"
    return header + "\n".join(selected), True


# A full-file read above this size (with no offset/limit) is refused in favour of
# a bounded head + a pointer to bio_inspect / bash, so one read can't blow up
# memory or the model's context (e.g. a 38 MB / multi-GB CSV).
_LARGE_TEXT_BYTES = 5 * 1024 * 1024  # 5 MB
_LARGE_TEXT_PREVIEW_LINES = 50


def _read_file(path: str, working_dir: str, offset: int = None, limit: int = None) -> tuple[str, bool]:
    _sess = _remote()
    if _sess is not None:
        return _remote_read_file(_sess, path, working_dir, offset, limit)
    resolved = _resolve(path, working_dir)
    if not resolved.exists():
        return f"File not found: {path}", False
    if not resolved.is_file():
        return f"Not a file: {path}", False

    # PDF → extract text
    if resolved.suffix.lower() == ".pdf":
        return _extract_pdf(resolved, offset, limit)

    # Excel → convert to CSV-like text
    if resolved.suffix.lower() in (".xlsx", ".xls", ".xlsm", ".ods"):
        return _extract_excel(resolved, offset, limit)

    # Word → extract text
    if resolved.suffix.lower() in (".docx", ".doc"):
        return _extract_docx(resolved, offset, limit)

    # Other binary → reject early with a clear message
    if _is_binary(resolved):
        size = resolved.stat().st_size
        return (
            f"Binary file: {resolved.name} ({size:,} bytes). "
            "Cannot read as text. Use a dedicated tool or convert it first.",
            False,
        )

    # Large text file with no explicit window → don't load the whole thing into
    # memory and the model's context. Return a bounded head + guidance so the
    # model reaches for the right tool (bio_inspect for tables, bash/pandas for
    # stats, or read_file with offset/limit to page). Paged reads (offset/limit
    # given) and small files fall through to the normal path unchanged.
    if offset is None and limit is None and resolved.stat().st_size > _LARGE_TEXT_BYTES:
        size = resolved.stat().st_size
        preview_lines: list[str] = []
        with open(resolved, "r", errors="replace") as fh:
            for i, line in enumerate(fh):
                if i >= _LARGE_TEXT_PREVIEW_LINES:
                    break
                preview_lines.append(line.rstrip("\n"))
        numbered = "\n".join(f"{i + 1}\t{ln}" for i, ln in enumerate(preview_lines))
        ext = resolved.suffix.lower()
        if ext in (".csv", ".tsv", ".tab", ".parquet", ".pq", ".jsonl", ".ndjson"):
            tip = f"This is a data table — call bio_inspect(path='{path}') for a schema-aware preview (shape, dtypes, summary)."
        else:
            tip = "Use bash (head/grep/awk) or read_file with offset/limit to page through it; do not read it whole."
        return (
            f"File: {path} ({size:,} bytes — too large to read whole; showing first "
            f"{len(preview_lines)} lines).\n{numbered}\n\n[LARGE FILE] {tip}",
            True,
        )

    try:
        lines = resolved.read_text(errors="replace").splitlines()
    except OSError as e:
        return str(e), False

    start = (offset - 1) if offset and offset > 0 else 0
    end = (start + limit) if limit else len(lines)
    selected = lines[start:end]

    numbered = "\n".join(f"{start + i + 1}\t{line}" for i, line in enumerate(selected))
    total = len(lines)
    header = f"File: {path} ({total} lines total)\n"
    return header + numbered, True


def _write_file(path: str, content: str, working_dir: str) -> tuple[str, bool]:
    _sess = _remote()
    if _sess is not None:
        return _remote_write_file(_sess, path, content, working_dir)
    resolved = _resolve(path, working_dir)
    size = len(content.encode("utf-8"))
    if size > _MAX_DOWNLOAD_BYTES:
        size_gb = size / (1024 ** 3)
        return (
            f"Refused to write {path} — content is {size_gb:.2f} GB, "
            f"exceeds the 1 GB limit.",
            False,
        )
    resolved.parent.mkdir(parents=True, exist_ok=True)
    resolved.write_text(content)
    lines = content.count("\n") + 1
    return f"Written {lines} lines to {path}", True


def _edit_file(
    path: str,
    old_string: str,
    new_string: str,
    working_dir: str,
    replace_all: bool = False,
) -> tuple[str, bool]:
    _sess = _remote()
    if _sess is not None:
        return _remote_edit_file(_sess, path, old_string, new_string, working_dir, replace_all)
    resolved = _resolve(path, working_dir)
    if not resolved.exists():
        return f"File not found: {path}", False
    if not resolved.is_file():
        return f"Not a file: {path}", False
    if old_string == "":
        return "old_string must not be empty. To create a file, use write_file.", False
    if old_string == new_string:
        return "old_string and new_string are identical — no change needed.", False
    if _is_binary(resolved):
        return f"Cannot edit binary file: {resolved.name}", False

    content = resolved.read_text(errors="replace")
    count = content.count(old_string)

    if count == 0:
        return f"String not found in {path}:\n{old_string[:200]}", False
    if count > 1 and not replace_all:
        return (
            f"old_string appears {count} times in {path} — "
            f"make it more specific to ensure uniqueness, or pass replace_all=true to replace every occurrence.",
            False,
        )

    new_content = content.replace(old_string, new_string)
    resolved.write_text(new_content)
    if replace_all and count > 1:
        return f"Edited {path} ({count} occurrences replaced)", True
    return f"Edited {path}", True


# Field-name aliases models emit for an edit's two halves. apply_patch accepts
# any of them so a slightly-off tool call still applies instead of hard-failing.
_EDIT_OLD_KEYS = ("old_string", "oldString", "old", "oldText", "old_text",
                  "search", "find", "before", "from")
_EDIT_NEW_KEYS = ("new_string", "newString", "new", "newText", "new_text",
                  "replace", "replacement", "after", "to")


def _first_key(d: dict, keys):
    for k in keys:
        if k in d and d[k] is not None:
            return d[k]
    return None


def _coerce_one_edit(e):
    """Normalize one edit-ish object to {old_string, new_string[, replace_all]},
    or None if it carries no old/new text. Tolerates a JSON-encoded string and
    alternate field names."""
    if isinstance(e, str):
        s = e.strip()
        if not s:
            return None
        try:
            e = json.loads(s)
        except Exception:
            return None
    if not isinstance(e, dict):
        return None
    old = _first_key(e, _EDIT_OLD_KEYS)
    new = _first_key(e, _EDIT_NEW_KEYS)
    if old is None and new is None:
        return None
    norm = {"old_string": "" if old is None else str(old),
            "new_string": "" if new is None else str(new)}
    ra = e.get("replace_all", e.get("replaceAll"))
    if ra is not None:
        norm["replace_all"] = bool(ra)
    return norm


def _coerce_edits(edits, extra=None):
    """Coerce the model's ``edits`` argument into a clean list[dict] of
    {old_string, new_string[, replace_all]}.

    Different models serialize the nested array in slightly different ways; this
    accepts the common ones instead of rejecting the call with "requires a
    non-empty edits list":
      - a JSON-encoded STRING of the array (``"[{...}]"``),
      - a single edit OBJECT (not wrapped in a list),
      - an index-keyed dict (``{"0": {...}, "1": {...}}``),
      - per-edit JSON strings, and alternate field names (old/new, search/
        replace, before/after, …),
      - a single top-level edit passed WITHOUT an ``edits`` wrapper (the model
        called apply_patch like edit_file), recovered from ``extra``.
    Returns [] when nothing usable is present.
    """
    if isinstance(edits, str):
        s = edits.strip()
        try:
            edits = json.loads(s) if s else None
        except Exception:
            edits = None
    if isinstance(edits, dict):
        # index-keyed dict → ordered list; otherwise a single edit object
        if edits and all(str(k).lstrip("-").isdigit() for k in edits):
            edits = [edits[k] for k in sorted(edits, key=lambda x: int(x))]
        else:
            edits = [edits]
    out: list[dict] = []
    if isinstance(edits, list):
        for e in edits:
            ne = _coerce_one_edit(e)
            if ne is not None:
                out.append(ne)
    # Fallbacks from top-level args when `edits` yielded nothing.
    if not out and isinstance(extra, dict) and extra:
        for k in ("patch", "changes", "edit", "hunks", "diffs"):
            if k in extra:
                return _coerce_edits(extra[k])
        ne = _coerce_one_edit(extra)
        if ne is not None:
            out.append(ne)
    return out


def _apply_patch(path: str, edits=None, working_dir: str = "", **extra) -> tuple[str, bool]:
    """Apply several old→new string replacements to one file, atomically.

    The file is written only after every edit has matched, so a failure
    mid-list leaves the file untouched. ``edits`` is normalized via
    :func:`_coerce_edits` first, so a slightly-malformed tool call (stringified
    array, single object, alt field names) still applies.
    """
    edits = _coerce_edits(edits, extra)
    _sess = _remote()
    if _sess is not None:
        return _remote_apply_patch(_sess, path, edits, working_dir)
    resolved = _resolve(path, working_dir)
    if not resolved.exists():
        return f"File not found: {path}", False
    if not resolved.is_file():
        return f"Not a file: {path}", False
    if not edits:
        return (
            "apply_patch got no usable edits. Pass `edits` as a JSON array of "
            '{old_string, new_string} objects — e.g. edits=[{"old_string": "foo", '
            '"new_string": "bar"}] — not a string or a single object. (file left UNCHANGED)'
        ), False
    if _is_binary(resolved):
        return f"Cannot edit binary file: {resolved.name}", False

    content = resolved.read_text(errors="replace")
    applied = 0
    notes: list[str] = []
    for i, e in enumerate(edits, 1):
        if not isinstance(e, dict):
            return f"Edit #{i} is not an object — expected {{old_string, new_string}}.", False
        old = e.get("old_string", "")
        new = e.get("new_string", "")
        replace_all = bool(e.get("replace_all", False))
        if old == "":
            return f"Edit #{i}: old_string must not be empty. (no changes written)", False
        if old == new:
            return f"Edit #{i}: old_string and new_string are identical. (no changes written)", False
        count = content.count(old)
        if count == 0:
            return (
                f"Edit #{i}: string not found:\n{old[:200]}\n"
                f"(applied {applied}/{len(edits)} so far — file left UNCHANGED)",
                False,
            )
        if count > 1 and not replace_all:
            return (
                f"Edit #{i}: old_string appears {count} times — make it unique or set "
                f"replace_all=true for this edit. (file left UNCHANGED)",
                False,
            )
        content = content.replace(old, new)
        applied += 1
        if replace_all and count > 1:
            notes.append(f"#{i}×{count}")

    resolved.write_text(content)
    suffix = f" ({', '.join(notes)} replace_all)" if notes else ""
    return f"Applied {applied} edit(s) to {path}{suffix}", True


# ---------------------------------------------------------------------------
# Task checklist (todo_write)
# ---------------------------------------------------------------------------

_TODO_STORE: dict[str, list[dict]] = {}
_VALID_TODO_STATUS = {"pending", "in_progress", "completed"}

# Alt field names / status spellings models use for a checklist item, so a
# slightly-off todo_write still populates the list instead of "no valid tasks".
_TODO_CONTENT_KEYS = ("content", "task", "text", "title", "step", "name",
                      "description", "todo", "item", "activity", "label")
_TODO_STATUS_KEYS = ("status", "state")
_TODO_STATUS_ALIASES = {
    "done": "completed", "complete": "completed", "completed": "completed",
    "finished": "completed", "closed": "completed", "resolved": "completed",
    "in_progress": "in_progress", "inprogress": "in_progress",
    "active": "in_progress", "doing": "in_progress", "working": "in_progress",
    "started": "in_progress", "current": "in_progress", "wip": "in_progress",
    "pending": "pending", "todo": "pending", "to_do": "pending",
    "open": "pending", "not_started": "pending", "queued": "pending",
    "planned": "pending", "new": "pending", "": "pending",
}


def _normalize_todo_status(s) -> str:
    key = str(s or "").strip().lower().replace(" ", "_").replace("-", "_")
    if key in _TODO_STATUS_ALIASES:
        return _TODO_STATUS_ALIASES[key]
    return key if key in _VALID_TODO_STATUS else "pending"


def _coerce_todos(todos):
    """Coerce the model's ``todos`` argument into a clean list[dict] of
    {content, status}. Same tolerance as :func:`_coerce_edits`: accepts a
    JSON-encoded string, a single object, an index-keyed dict, plain-string items
    (``["step 1", "step 2"]`` → each becomes content with status 'pending'),
    per-item JSON strings, alternate content keys (task/text/title/step/…) and
    status spellings (done→completed, doing→in_progress, todo→pending, …).
    Returns [] when nothing usable is present."""
    if isinstance(todos, str):
        s = todos.strip()
        try:
            todos = json.loads(s) if s else None
        except Exception:
            todos = None
    if isinstance(todos, dict):
        if todos and all(str(k).lstrip("-").isdigit() for k in todos):
            todos = [todos[k] for k in sorted(todos, key=lambda x: int(x))]
        else:
            todos = [todos]
    out: list[dict] = []
    if isinstance(todos, list):
        for t in todos:
            if isinstance(t, str):
                st = t.strip()
                if not st:
                    continue
                # A string item may itself be a JSON object; else it's content.
                if st[:1] in "{[":
                    try:
                        t = json.loads(st)
                    except Exception:
                        out.append({"content": st, "status": "pending"})
                        continue
                else:
                    out.append({"content": st, "status": "pending"})
                    continue
            if not isinstance(t, dict):
                continue
            content = _first_key(t, _TODO_CONTENT_KEYS)
            if content is None:
                continue
            content = str(content).strip()
            if not content:
                continue
            status = _normalize_todo_status(_first_key(t, _TODO_STATUS_KEYS))
            out.append({"content": content, "status": status})
    return out


def get_todos(working_dir: str) -> list[dict]:
    """Return the current task list for a working dir (used by the web UI)."""
    return list(_TODO_STORE.get(working_dir, []))


def _todo_write(todos=None, working_dir: str = "", **extra) -> tuple[str, bool]:
    norm = _coerce_todos(todos)
    # Fallbacks when `todos` yielded nothing: an alt container key, or a single
    # top-level {content, status} passed without a `todos` wrapper.
    if not norm and extra:
        for k in ("items", "tasks", "list", "todo", "checklist", "steps"):
            if k in extra:
                norm = _coerce_todos(extra[k])
                if norm:
                    break
        if not norm:
            norm = _coerce_todos(extra)

    if not norm:
        return "todo_write received no valid tasks (each needs `content` and `status`).", False

    _TODO_STORE[working_dir] = norm
    try:
        from . import display
        display.print_todos(norm)
    except Exception:
        pass

    done = sum(1 for t in norm if t["status"] == "completed")
    in_prog = [t["content"] for t in norm if t["status"] == "in_progress"]
    summary = f"Task list updated — {done}/{len(norm)} completed."
    if in_prog:
        summary += f" Now working on: {in_prog[0]}"
    return summary, True


# ---------------------------------------------------------------------------
# Background processes (run_background / check_process / stop_process)
# ---------------------------------------------------------------------------

_BG_PROCS: dict[str, dict] = {}
_bg_counter = 0
_bg_lock = threading.Lock()


def _remote_run_background(sess, command: str, working_dir: str, cwd: str = None) -> tuple[str, bool]:
    import shlex
    global _bg_counter
    run_dir = cwd or working_dir
    with _bg_lock:
        _bg_counter += 1
        proc_id = f"bg{_bg_counter}"
    log_path = _remote_resolve(f".ots_bg_{proc_id}.log", run_dir)
    launch = (f"nohup sh -c {shlex.quote(command)} > {shlex.quote(log_path)} 2>&1 "
              f"< /dev/null & echo $!")
    out, err, rc = sess.run(launch, cwd=run_dir, timeout=60)
    pid = out.strip().splitlines()[-1] if out.strip() else ""
    if rc != 0 or not pid.isdigit():
        return f"Failed to start remote background process: {(err or out).strip()}", False
    with _bg_lock:
        _BG_PROCS[proc_id] = {
            "remote": True,
            "sess": sess,
            "pid": pid,
            "command": command,
            "log": log_path,
            "started": _time.time(),
            "cwd": run_dir,
        }
    return (
        f"Started background process {proc_id} (pid {pid}) on remote: {command}\n"
        f"Use check_process(id=\"{proc_id}\") to read output, "
        f"stop_process(id=\"{proc_id}\") to stop it.",
        True,
    )


def _run_background(command: str, working_dir: str, cwd: str = None) -> tuple[str, bool]:
    global _bg_counter
    _sess = _remote()
    if _sess is not None:
        return _remote_run_background(_sess, command, working_dir, cwd)
    run_dir = cwd or working_dir
    try:
        log_fd, log_path = tempfile.mkstemp(prefix="ots_bg_", suffix=".log")
        log_file = os.fdopen(log_fd, "w")
    except OSError as e:
        return f"Could not allocate a log file: {e}", False

    env = {k: v for k, v in os.environ.items() if k != "VIRTUAL_ENV"}
    try:
        proc = subprocess.Popen(
            command,
            shell=True,
            cwd=run_dir,
            stdout=log_file,
            stderr=subprocess.STDOUT,
            text=True,
            env=env,
            start_new_session=True,  # own process group → killable as a unit
        )
    except Exception as e:
        try:
            log_file.close()
        except Exception:
            pass
        return f"Failed to start background process: {e}", False

    with _bg_lock:
        _bg_counter += 1
        proc_id = f"bg{_bg_counter}"
        _BG_PROCS[proc_id] = {
            "proc": proc,
            "command": command,
            "log": log_path,
            "file": log_file,
            "started": _time.time(),
            "cwd": run_dir,
        }
    return (
        f"Started background process {proc_id} (pid {proc.pid}): {command}\n"
        f"Use check_process(id=\"{proc_id}\") to read output, "
        f"stop_process(id=\"{proc_id}\") to stop it.",
        True,
    )


def _read_bg_log(info: dict, tail_lines: int) -> str:
    try:
        info["file"].flush()
    except Exception:
        pass
    try:
        with open(info["log"], errors="replace") as fh:
            lines = fh.read().splitlines()
    except OSError:
        return "(could not read output)"
    if tail_lines and len(lines) > tail_lines:
        return f"[last {tail_lines} of {len(lines)} lines]\n" + "\n".join(lines[-tail_lines:])
    return "\n".join(lines) if lines else "(no output yet)"


def _remote_bg_alive(info: dict) -> bool:
    return info["sess"].run(f"kill -0 {info['pid']} 2>/dev/null", None, timeout=30)[2] == 0


def _read_remote_bg_log(info: dict, tail_lines: int) -> str:
    import shlex
    n = tail_lines or 200
    out, _, rc = info["sess"].run(
        f"tail -n {n} {shlex.quote(info['log'])} 2>/dev/null", None, timeout=60)
    if rc != 0:
        return "(could not read output)"
    return out.rstrip("\n") if out.strip() else "(no output yet)"


def _check_process(working_dir: str, id: str = None, tail_lines: int = 50) -> tuple[str, bool]:
    if not id:
        with _bg_lock:
            items = list(_BG_PROCS.items())
        if not items:
            return "No background processes have been started.", True
        lines = ["Background processes:"]
        for pid_id, info in items:
            if info.get("remote"):
                status = "running" if _remote_bg_alive(info) else "exited"
            else:
                rc = info["proc"].poll()
                status = "running" if rc is None else f"exited ({rc})"
            lines.append(f"  {pid_id}: {status} — {info['command'][:70]}")
        return "\n".join(lines), True

    info = _BG_PROCS.get(id)
    if not info:
        return f"No background process with id '{id}'. Use check_process() to list them.", False
    if info.get("remote"):
        status = "running" if _remote_bg_alive(info) else "exited"
        elapsed = int(_time.time() - info["started"])
        header = f"Process {id} — {status} (running {elapsed}s): {info['command'][:80]}\n"
        return header + _read_remote_bg_log(info, tail_lines), True
    rc = info["proc"].poll()
    status = "running" if rc is None else f"exited with code {rc}"
    elapsed = int(_time.time() - info["started"])
    header = f"Process {id} — {status} (running {elapsed}s): {info['command'][:80]}\n"
    return header + _read_bg_log(info, tail_lines), True


def running_background_processes() -> list[dict]:
    """Public: background jobs started this session that are STILL running.

    Returns ``[{"id", "command", "elapsed"}, …]`` (elapsed in whole seconds).
    Used by the council loop to refuse a "task complete" claim while a job the
    worker launched is still in flight — otherwise a turn-based loop would end
    with the real deliverable never produced. Never raises."""
    out: list[dict] = []
    with _bg_lock:
        items = list(_BG_PROCS.items())
    for proc_id, info in items:
        try:
            if info.get("remote"):
                alive = _remote_bg_alive(info)
            else:
                alive = info["proc"].poll() is None
        except Exception:
            alive = False
        if alive:
            out.append({
                "id": proc_id,
                "command": info.get("command", ""),
                "elapsed": int(_time.time() - info.get("started", _time.time())),
            })
    return out


def _stop_process(id: str, working_dir: str) -> tuple[str, bool]:
    info = _BG_PROCS.get(id)
    if not info:
        return f"No background process with id '{id}'.", False
    if info.get("remote"):
        if not _remote_bg_alive(info):
            return f"Process {id} already exited.", True
        pid = info["pid"]
        info["sess"].run(f"kill {pid} 2>/dev/null; sleep 1; kill -9 {pid} 2>/dev/null || true",
                         None, timeout=30)
        return f"Stopped background process {id}.", True
    proc = info["proc"]
    if proc.poll() is not None:
        return f"Process {id} already exited (code {proc.poll()}).", True
    _terminate_proc(proc)
    try:
        info["file"].close()
    except Exception:
        pass
    return f"Stopped background process {id}.", True


def _terminate_proc(proc: subprocess.Popen) -> None:
    """SIGTERM the process group, escalate to SIGKILL if it doesn't exit."""
    try:
        os.killpg(os.getpgid(proc.pid), signal.SIGTERM)
    except Exception:
        try:
            proc.terminate()
        except Exception:
            return
    try:
        proc.wait(timeout=5)
    except Exception:
        try:
            os.killpg(os.getpgid(proc.pid), signal.SIGKILL)
        except Exception:
            try:
                proc.kill()
            except Exception:
                pass


def _shutdown_bg_procs() -> None:
    """Kill any still-running local background processes on interpreter exit.

    Remote background jobs are detached (nohup) and intentionally survive; the
    agent stops them explicitly via stop_process.
    """
    for info in list(_BG_PROCS.values()):
        if info.get("remote"):
            continue
        try:
            if info["proc"].poll() is None:
                _terminate_proc(info["proc"])
        except Exception:
            pass


atexit.register(_shutdown_bg_procs)


# ---------------------------------------------------------------------------
# ask_user
# ---------------------------------------------------------------------------

def _ask_user(question: str, working_dir: str, options: list = None) -> tuple[str, bool]:
    opts = [str(o) for o in options] if isinstance(options, list) else None
    try:
        from . import display
        answer, answered = display.ask_user_question(question, opts)
    except Exception:
        answered, answer = False, ""
    if not answered:
        return (
            "No answer was received — the user is unavailable or this is a "
            "non-interactive/autonomous run. Proceed with your best judgement "
            "based on the task and sensible defaults; do not ask again.",
            True,
        )
    return f"The user answered: {answer}", True


def _remember(content: str, working_dir: str, tags=None) -> tuple[str, bool]:
    """Persist an agent-authored insight to this project's memory."""
    # Lazy import — agent.py imports tools.py at module load, so importing it
    # at the top here would be a circular import.
    from .agent import save_memory_insight
    ok = save_memory_insight(working_dir, content, tags)
    if not ok:
        return "Could not save the insight (empty content or write error).", False
    preview = content.strip().replace("\n", " ")
    if len(preview) > 80:
        preview = preview[:77] + "…"
    return f"Saved to memory: {preview}", True


def _forget(content: str, working_dir: str) -> tuple[str, bool]:
    """Remove stale insight(s) from this project's memory matching ``content``."""
    from .agent import delete_memory_insight
    removed = delete_memory_insight(working_dir, content)
    if not removed:
        return "No matching insight found in memory; nothing removed.", True
    previews = []
    for r in removed:
        r = r.replace("\n", " ")
        previews.append(r if len(r) <= 80 else r[:77] + "…")
    return "Removed from memory:\n- " + "\n- ".join(previews), True


# Commands that essentially never exit on their own — dev servers, watchers,
# REPLs, log tails. Running these through bash blocks the whole agent (and, in a
# team run, the whole lab) until the timeout fires. We intercept them up front
# and redirect to run_background instead of executing. Kept tight to
# unambiguous patterns so ordinary one-shot commands are never caught; the model
# can still force a foreground run by prefixing `OTS_FOREGROUND=1 ` (e.g. a
# server invoked only for `--help`).
_BLOCKING_CMD_PATTERNS = (
    r"\bflask\s+run\b",
    r"\b(uvicorn|gunicorn|hypercorn|daphne|waitress-serve)\b",
    r"\bmanage\.py\s+runserver\b",
    r"\bstreamlit\s+run\b",
    r"\bjupyter\s+(notebook|lab)\b",
    r"\bhttp\.server\b",
    r"\b(nodemon|webpack-dev-server|next\s+dev|vite(\s+dev)?)\b",
    r"\b(npm|yarn|pnpm)\s+(run\s+)?(dev|start|serve)\b",
    r"\btail\s+-[a-zA-Z]*f",
    r"--watch\b",
)


def _looks_blocking(command: str) -> bool:
    import re
    if "OTS_FOREGROUND=1" in command:
        return False
    return any(re.search(p, command) for p in _BLOCKING_CMD_PATTERNS)


class _Completed:
    """Minimal stand-in for ``subprocess.CompletedProcess``."""
    __slots__ = ("stdout", "stderr", "returncode")

    def __init__(self, stdout: str, stderr: str, returncode: int):
        self.stdout, self.stderr, self.returncode = stdout, stderr, returncode


def _run_interruptible(command: str, working_dir: str, timeout: int,
                       env: dict) -> "_Completed | None":
    """Run ``command`` like ``subprocess.run``, but abortable.

    ``subprocess.run`` blocks until the command exits or its timeout fires, which
    is why pressing Stop during a long build used to do nothing for minutes. Here
    the process gets its OWN process group and we poll while waiting, so a stop
    can kill the whole tree (the shell AND its children) at once.

    Returns None when the user stopped the session; raises ``TimeoutExpired`` on
    the normal timeout so the caller's existing handling is unchanged.
    """
    from . import interrupt
    proc = subprocess.Popen(
        command, shell=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE,
        text=True, cwd=working_dir, env=env, start_new_session=True,
    )
    deadline = _time.monotonic() + timeout
    while True:
        try:
            stdout, stderr = proc.communicate(timeout=0.25)
            return _Completed(stdout, stderr, proc.returncode)
        except subprocess.TimeoutExpired:
            if interrupt.should_stop():
                _terminate_proc(proc)
                try:            # drain whatever it managed to print
                    proc.communicate(timeout=2)
                except Exception:
                    pass
                return None
            if _time.monotonic() >= deadline:
                _terminate_proc(proc)
                try:
                    proc.communicate(timeout=2)
                except Exception:
                    pass
                raise subprocess.TimeoutExpired(command, timeout)


def _bash(command: str, working_dir: str, timeout: int = 300) -> tuple[str, bool]:
    # Redirect commands that never return on their own to run_background, so a
    # single blocking call can't freeze the agent until the timeout fires.
    if _looks_blocking(command):
        return (
            f"This looks like a long-running / blocking process (dev server, watcher, "
            f"or log tail) that won't exit on its own — running it through bash would "
            f"block everything until the {timeout}s timeout.\n"
            f"Start it with run_background instead: "
            f"run_background(command={command!r}), then poll it with check_process and "
            f"stop it with stop_process when done.\n"
            f"If you genuinely need it in the foreground (e.g. a `--help`/version check), "
            f"prefix the command with `OTS_FOREGROUND=1 ` to bypass this guard.",
            False,
        )
    _sess = _remote()
    if _sess is not None:
        return _remote_bash(_sess, command, working_dir, timeout)
    # Unset VIRTUAL_ENV so uv doesn't emit a mismatch warning when the conda/system
    # venv doesn't match the project's .venv.
    env = {k: v for k, v in os.environ.items() if k != "VIRTUAL_ENV"}
    try:
        result = _run_interruptible(command, working_dir, timeout, env)
        if result is None:
            # User pressed Stop while this command was running. The process (and
            # everything it spawned) has been killed; report it as the tool result
            # so the agent's history says what happened instead of showing a
            # truncated command that appears to have failed on its own.
            return ("Command was KILLED because the user stopped the session. "
                    "It did not finish, so its effects may be partial — check the "
                    "actual state before assuming anything about it.", False)
        stdout = result.stdout or ""
        stderr = result.stderr or ""
        # Label streams when both have content so the model can tell them apart.
        # Most successful commands write only to stdout — keep that path lean.
        if stdout and stderr:
            output = f"--- stdout ---\n{stdout}\n--- stderr ---\n{stderr}"
        else:
            output = stdout or stderr
        if not output:
            output = f"(exit code {result.returncode})"
        elif result.returncode != 0:
            output += f"\n(exit code {result.returncode})"
        # Truncate very long outputs — keep tail-heavy since errors appear at the end
        if len(output) > 8000:
            output = output[:2000] + "\n\n... [output truncated] ...\n\n" + output[-5000:]
        return output, result.returncode == 0
    except subprocess.TimeoutExpired:
        return (
            f"Command timed out after {timeout}s and was killed. If this job is "
            f"genuinely long-running (training, simulation, big data job, a server), "
            f"do NOT just re-run with a bigger timeout — that blocks the agent the "
            f"whole time. Re-launch it with run_background(command=...) and poll it "
            f"with check_process. Only raise the bash timeout for a bounded one-off "
            f"(a big install or full test suite) that you expect to finish.",
            False,
        )


def _find_codag() -> str | None:
    """Locate the codag binary. Returns absolute path, or None if not installed."""
    found = shutil.which("codag")
    if found:
        return found
    # install.sh default location
    fallback = Path.home() / ".local" / "bin" / "codag"
    return str(fallback) if fallback.is_file() and os.access(fallback, os.X_OK) else None


# Module-level flag so we only attempt auto-install once per process. Without
# this a failing install (no network, locked-down system) would re-attempt on
# every compress_log call.
_codag_autoinstall_attempted = False


def _attempt_codag_install() -> bool:
    """One-shot best-effort install of the codag CLI for users who got
    octoslave via the platform installer (DMG/EXE/AppImage) — they never ran
    scripts/install.sh and so don't have codag yet. Returns True if codag is
    found on PATH afterward.

    Opt out via OCTOSLAVE_NO_CODAG_AUTOINSTALL=1. Requires curl + sh (Unix);
    on Windows or stripped-down environments, this is a no-op and the caller
    gets the standard "not installed" error.
    """
    global _codag_autoinstall_attempted
    if _codag_autoinstall_attempted:
        return _find_codag() is not None
    _codag_autoinstall_attempted = True

    if os.environ.get("OCTOSLAVE_NO_CODAG_AUTOINSTALL") == "1":
        return False
    if not shutil.which("curl") or not shutil.which("sh"):
        return False

    try:
        from . import display
        display.print_info(
            "compress_log: codag CLI not installed — fetching it now (one-time, ~5s)…"
        )
    except Exception:
        pass

    try:
        result = subprocess.run(
            "curl -fsSL https://codag.ai/install.sh | sh",
            shell=True,
            capture_output=True,
            text=True,
            timeout=120,
        )
    except subprocess.TimeoutExpired:
        try:
            from . import display
            display.print_info(
                "compress_log: codag install timed out — falling back to raw output. "
                "Re-run with OCTOSLAVE_NO_CODAG_AUTOINSTALL=1 to silence this on future calls."
            )
        except Exception:
            pass
        return False
    except Exception:
        return False

    ok = result.returncode == 0 and _find_codag() is not None
    try:
        from . import display
        if ok:
            display.print_info("compress_log: codag installed.")
        else:
            tail = (result.stderr or result.stdout or "").strip().splitlines()
            tail_line = tail[-1] if tail else "(no output)"
            display.print_info(
                f"compress_log: codag install failed ({tail_line}). "
                "Install manually with `curl -fsSL https://codag.ai/install.sh | sh`."
            )
    except Exception:
        pass
    return ok


# Codag's free `compact` mode has a server-side cap of 5,000 input lines per
# request. To handle real ML/CI logs we chunk above that. Headroom of 500.
_CODAG_MAX_LINES_PER_CHUNK = 4500
# Hard ceiling on chunk count to avoid burning many API calls on truly huge
# files. 10 chunks * 4500 lines = ~45K input lines covered with full fidelity.
_CODAG_MAX_CHUNKS = 10


def _run_codag(codag: str, argv_prefix: list[str], stdin_text: str, cwd: str, timeout: int) -> tuple[str, str, int]:
    """Pipe stdin_text to `codag wrap ...` and return (stdout, stderr, returncode)."""
    env = {k: v for k, v in os.environ.items() if k != "VIRTUAL_ENV"}
    try:
        result = subprocess.run(
            argv_prefix,
            input=stdin_text,
            capture_output=True,
            text=True,
            cwd=cwd,
            timeout=timeout,
            env=env,
        )
        return result.stdout or "", result.stderr or "", result.returncode
    except subprocess.TimeoutExpired:
        return "", f"codag timed out after {timeout}s", 124
    except FileNotFoundError:
        return "", "codag binary disappeared between lookup and exec. Reinstall codag.", 127


def _collect_lines(
    working_dir: str,
    command: str | None,
    path: str | None,
    timeout: int,
) -> tuple[list[str] | None, str | None]:
    """Resolve input into a list of lines. Returns (lines, error_message)."""
    if path:
        resolved = _resolve(path, working_dir)
        if not resolved.exists():
            return None, f"compress_log: file not found: {resolved}"
        if not resolved.is_file():
            return None, f"compress_log: not a regular file: {resolved}"
        try:
            text = resolved.read_text(encoding="utf-8", errors="replace")
        except OSError as e:
            return None, f"compress_log: could not read {resolved}: {e}"
        return text.splitlines(), None

    # command mode — run it ourselves so we control the input shape sent to codag
    env = {k: v for k, v in os.environ.items() if k != "VIRTUAL_ENV"}
    try:
        result = subprocess.run(
            ["sh", "-c", command],
            capture_output=True,
            text=True,
            cwd=working_dir,
            timeout=timeout,
            env=env,
        )
    except subprocess.TimeoutExpired:
        return None, f"compress_log: command timed out after {timeout}s"
    raw = (result.stdout or "") + ("\n" + result.stderr if result.stderr else "")
    return raw.splitlines(), None


def _compress_log(
    working_dir: str,
    command: str = None,
    path: str = None,
    service: str = None,
    mode: str = "compact",
    timeout: int = 600,
) -> tuple[str, str]:
    """Compress a log/command-output via the `codag` CLI.

    Free `compact` mode is capped at 5,000 input lines per request server-side;
    larger inputs are split into ≤4,500-line chunks and compressed sequentially,
    with each chunk's summary stitched together. Up to 10 chunks (~45K lines)
    are processed; beyond that, the input is head+tail truncated to fit and a
    notice is appended to the result.
    """
    if (command and path) or (not command and not path):
        return (
            "compress_log requires exactly one of `command` or `path` "
            "(got both, or neither). Pass `path` to compress a file, or `command` "
            "to wrap and compress a shell command's output.",
            False,
        )

    codag = _find_codag()
    if codag is None:
        # First-call auto-install for platform-installer users who never ran
        # scripts/install.sh. Best-effort; opt out via OCTOSLAVE_NO_CODAG_AUTOINSTALL=1.
        if _attempt_codag_install():
            codag = _find_codag()
    if codag is None:
        return (
            "codag CLI not installed and auto-install failed. Install manually with: "
            "`curl -fsSL https://codag.ai/install.sh | sh` "
            "(installs to ~/.local/bin/codag — make sure that dir is on PATH). "
            "Then retry compress_log.",
            False,
        )

    # Also register codag's richer MCP server (wrap, compact, health, plus
    # tail_* log-tailers) on first successful use, so future runs of any role
    # can use them directly. Tail tools whose CLI is missing (docker/kubectl/
    # aws/gh/vercel) are pruned at load — see mcp_registry.unsupported_tools.
    # Idempotent; opt out via OCTOSLAVE_NO_CODAG_MCP_REGISTER=1.
    try:
        from .config import ensure_codag_mcp_registered
        if ensure_codag_mcp_registered():
            from . import display
            display.print_info(
                "compress_log: registered the codag MCP server "
                "(exposes wrap / compact / health plus tail_* tools for any "
                "installed log CLIs). Restart octoslave to load it."
            )
    except Exception:
        pass

    if mode not in ("compact", "capsule", "parse"):
        return f"Invalid mode '{mode}'. Use 'compact' (default, no auth) or 'capsule' (JSON, needs auth).", False

    lines, err = _collect_lines(working_dir, command, path, timeout)
    if err is not None:
        return err, False
    if not lines:
        return "(input was empty — nothing to compress)", True

    argv_prefix = [codag, "wrap", "--mode", mode]
    if service:
        argv_prefix += ["--service", service]
    # codag reads from stdin when no `-- cmd` is provided.

    # Single-shot path
    if len(lines) <= _CODAG_MAX_LINES_PER_CHUNK:
        stdout, stderr, rc = _run_codag(codag, argv_prefix, "\n".join(lines), working_dir, timeout)
        if rc != 0:
            return (
                f"codag failed (exit {rc}):\n{(stderr or stdout).strip()}\n"
                f"Tip: `capsule`/`parse` modes need `codag auth login`; "
                f"`compact` is free.",
                False,
            )
        summary_line = next((ln.strip() for ln in stderr.splitlines() if ln.startswith("[codag]")), "")
        out = stdout.strip()
        return (f"{summary_line}\n{out}" if summary_line else out) or "(empty output)", True

    # Chunked path: decide whether all lines fit within MAX_CHUNKS,
    # otherwise head+tail truncate so the most valuable lines survive.
    note = ""
    max_lines_total = _CODAG_MAX_LINES_PER_CHUNK * _CODAG_MAX_CHUNKS
    if len(lines) > max_lines_total:
        head_count = max_lines_total // 2
        tail_count = max_lines_total - head_count
        truncated_count = len(lines) - max_lines_total
        lines = (
            lines[:head_count]
            + [f"... [{truncated_count:,} middle lines omitted by compress_log — exceeds chunk budget] ..."]
            + lines[-tail_count:]
        )
        note = (
            f"\n[compress_log: input was {len(lines) + truncated_count - 1:,} lines; "
            f"head {head_count:,} + tail {tail_count:,} were compressed, middle dropped]"
        )

    # Split into chunks
    chunks: list[list[str]] = [
        lines[i : i + _CODAG_MAX_LINES_PER_CHUNK]
        for i in range(0, len(lines), _CODAG_MAX_LINES_PER_CHUNK)
    ]
    total = len(chunks)
    pieces: list[str] = []
    for idx, chunk in enumerate(chunks, 1):
        stdout, stderr, rc = _run_codag(codag, argv_prefix, "\n".join(chunk), working_dir, timeout)
        if rc != 0:
            return (
                f"codag failed on chunk {idx}/{total} (exit {rc}):\n{(stderr or stdout).strip()}",
                False,
            )
        summary_line = next((ln.strip() for ln in stderr.splitlines() if ln.startswith("[codag]")), "")
        body = stdout.strip()
        header = f"=== chunk {idx}/{total} ({len(chunk):,} lines)"
        if summary_line:
            header += f" — {summary_line[len('[codag]'):].strip()}"
        header += " ==="
        pieces.append(f"{header}\n{body}")

    return ("\n\n".join(pieces) + note).strip() or "(empty output)", True


_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".tiff", ".tif", ".bmp", ".gif", ".webp"}


def _image_ocr(
    path: str,
    working_dir: str,
    lang: str = "eng",
    preprocess: bool = False,
    psm: int = 3,
) -> tuple[str, bool]:
    try:
        from PIL import Image
    except ImportError:
        return "Pillow is not installed. Run: pip install pillow", False
    try:
        import pytesseract  # type: ignore
    except ImportError:
        return "pytesseract is not installed. Run: pip install pytesseract", False

    resolved = _resolve(path, working_dir)
    if not resolved.exists():
        return f"File not found: {path}", False
    if not resolved.is_file():
        return f"Not a file: {path}", False
    if resolved.suffix.lower() not in _IMAGE_EXTS:
        return (
            f"Unsupported image format: {resolved.suffix!r}. "
            f"Supported: {sorted(_IMAGE_EXTS)}. For PDFs use pdf_ocr.",
            False,
        )

    try:
        img = Image.open(str(resolved))
        img.load()
    except Exception as e:
        return f"Could not open image: {e}", False

    width, height = img.size

    if preprocess:
        img = img.convert("L")
        img = img.point(lambda x: 0 if x < 140 else 255, "1")

    config = f"--psm {int(psm)}"
    try:
        text = pytesseract.image_to_string(img, lang=lang, config=config)
    except pytesseract.TesseractNotFoundError:
        return (
            "tesseract binary not found on PATH. Install: "
            "macOS `brew install tesseract`, Debian `apt install tesseract-ocr`.",
            False,
        )
    except Exception as e:
        return f"OCR error: {e}", False

    header = f"OCR: {resolved.name} ({width}x{height}, lang={lang}, psm={psm})\n\n"
    if not text.strip():
        hint = "" if preprocess else " Try preprocess=true for noisy/low-contrast images."
        return header + f"(no text extracted —{hint} image may contain no readable text)", True
    return header + text, True


# ---------------------------------------------------------------------------
# view_image — hand the model the actual pixels
# ---------------------------------------------------------------------------
# The Chat Completions API this agent speaks requires a `role: "tool"` message's
# content to be a plain string, so an image cannot travel back as a tool result.
# Instead the tool returns a short text ack and parks the encoded image on this
# thread's pending queue; the agent loop drains the queue after the turn's tool
# results and appends the pixels as a separate `role: "user"` message (see
# agent.drain_image_attachments). Queue is thread-local, like the rest of _EXEC,
# so parallel agents never cross-contaminate.

# A 300-DPI matplotlib PNG is ~2 MB and every provider downsamples it anyway.
_VIEW_IMAGE_MAX_EDGE = 1280
_VIEW_IMAGE_MAX_BYTES = 1_500_000


def set_vision_support(check) -> None:
    """Declare whether THIS thread's model can accept image input.

    ``check`` is a bool, a zero-arg callable returning bool (evaluated lazily, so
    a live capability probe only costs a request if view_image is actually
    called), or None for "unknown" — unknown lets the image through and relies on
    the agent's 400-recovery to strip it if the endpoint refuses."""
    _EXEC.vision_check = check


def current_vision_support():
    """The raw check installed for this thread (bool | callable | None). Lets a
    nested run (Lab/Science specialist) save and restore the ambient setting."""
    return getattr(_EXEC, "vision_check", None)


def vision_supported() -> bool | None:
    check = getattr(_EXEC, "vision_check", None)
    if check is None or isinstance(check, bool):
        return check
    try:
        return bool(check())
    except Exception:
        return None


def take_image_attachments() -> list[dict]:
    """Drain this thread's pending image attachments (oldest first)."""
    pending = getattr(_EXEC, "image_attachments", None)
    if not pending:
        return []
    _EXEC.image_attachments = []
    return pending


def _queue_image_attachment(entry: dict) -> None:
    pending = getattr(_EXEC, "image_attachments", None)
    if pending is None:
        pending = []
        _EXEC.image_attachments = pending
    pending.append(entry)


def _encode_for_vision(img, max_bytes: int) -> tuple[str, bytes]:
    """Serialise a PIL image to (mime, bytes) under ``max_bytes``.

    PNG first — plots, gels and structures are line art, where PNG is both
    smaller and sharper than JPEG. Only photographic content overflows the
    budget, and that is exactly the content JPEG handles well."""
    import io
    buf = io.BytesIO()
    img.save(buf, format="PNG", optimize=True)
    if buf.tell() <= max_bytes:
        return "image/png", buf.getvalue()

    flat = img
    if flat.mode not in ("RGB", "L"):
        from PIL import Image
        bg = Image.new("RGB", flat.size, (255, 255, 255))
        bg.paste(flat, mask=flat.split()[-1] if flat.mode in ("RGBA", "LA") else None)
        flat = bg
    for quality in (85, 70, 55):
        buf = io.BytesIO()
        flat.save(buf, format="JPEG", quality=quality, optimize=True)
        if buf.tell() <= max_bytes:
            return "image/jpeg", buf.getvalue()
    # Still oversized (very large photo): halve it and take whatever JPEG gives.
    flat = flat.resize((max(1, flat.width // 2), max(1, flat.height // 2)))
    buf = io.BytesIO()
    flat.save(buf, format="JPEG", quality=70, optimize=True)
    return "image/jpeg", buf.getvalue()


# Formats every vision endpoint accepts verbatim. TIFF/BMP must be converted
# first, which needs Pillow.
_VIEW_IMAGE_RAW_MIMES = {
    ".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
    ".gif": "image/gif", ".webp": "image/webp",
}


def _view_image_raw(resolved: Path, display_path: str, note: str) -> tuple[str, bool]:
    """Attach an image without Pillow: only viable if it is already small enough
    and in a format the providers take as-is."""
    mime = _VIEW_IMAGE_RAW_MIMES.get(resolved.suffix.lower())
    if mime is None:
        return (
            f"Cannot view a {resolved.suffix} image without Pillow (it needs "
            f"converting first). Run `pip install pillow`, or convert the file "
            f"to PNG and view that.",
            False,
        )
    try:
        blob = resolved.read_bytes()
    except OSError as e:
        return f"Could not read image: {e}", False
    if len(blob) > _VIEW_IMAGE_MAX_BYTES:
        return (
            f"Image is {len(blob) // 1024} KB, over the "
            f"{_VIEW_IMAGE_MAX_BYTES // 1024} KB limit, and Pillow is not "
            f"installed to downscale it. Run `pip install pillow`, or save the "
            f"figure at a lower dpi and view that.",
            False,
        )
    _queue_image_attachment({
        "label": display_path,
        "mime": mime,
        "b64": base64.b64encode(blob).decode("ascii"),
        "note": (note or "").strip(),
    })
    return (
        f"Attached {Path(display_path).name} — {len(blob) // 1024} KB. The image "
        f"itself follows in the next message; read it there.",
        True,
    )


def _view_image(
    path: str,
    working_dir: str,
    note: str = "",
    label: str = "",
) -> tuple[str, bool]:
    """Show an image to the model itself (plots, structures, gels, micrographs).

    Returns a text ack; the pixels are queued for the agent loop to attach."""
    supported = vision_supported()
    if supported is False:
        return (
            "This model cannot accept image input, so the picture was NOT shown "
            "to you — do not describe it as if you had seen it. Use image_ocr to "
            "read any text/labels in it, inspect the underlying data file, or "
            "switch to a vision-capable model. (Override the detection with "
            "`model_vision` in config.json if this model really does see images.)",
            False,
        )

    resolved = _resolve(path, working_dir)
    if not resolved.exists():
        return f"File not found: {path}", False
    if not resolved.is_file():
        return f"Not a file: {path}", False
    if resolved.suffix.lower() not in _IMAGE_EXTS:
        return (
            f"Unsupported image format: {resolved.suffix!r}. "
            f"Supported: {sorted(_IMAGE_EXTS)}. For PDFs, render a page to PNG "
            f"first (or use pdf_ocr for text).",
            False,
        )

    try:
        from PIL import Image
    except ImportError:
        # Pillow is an optional extra, so fall back to sending the file as-is.
        # That covers most real figures (a default-DPI matplotlib PNG is well
        # under the cap); only oversized or non-web formats actually need the
        # resizer, and those say so instead of failing vaguely.
        return _view_image_raw(resolved, label or path, note)

    try:
        img = Image.open(str(resolved))
        img.load()
    except Exception as e:
        return f"Could not open image: {e}", False

    width, height = img.size
    if img.mode not in ("RGB", "RGBA", "L"):
        img = img.convert("RGB")
    scale = _VIEW_IMAGE_MAX_EDGE / max(width, height)
    if scale < 1:
        img = img.resize((max(1, round(width * scale)), max(1, round(height * scale))),
                         Image.LANCZOS)

    try:
        mime, blob = _encode_for_vision(img, _VIEW_IMAGE_MAX_BYTES)
    except Exception as e:
        return f"Could not encode image for viewing: {e}", False

    display_path = label or path
    _queue_image_attachment({
        "label": display_path,
        "mime": mime,
        "b64": base64.b64encode(blob).decode("ascii"),
        "note": (note or "").strip(),
    })

    sent = f"{img.width}x{img.height}" if scale < 1 else f"{width}x{height}"
    scaled = f" (downscaled from {width}x{height})" if scale < 1 else ""
    return (
        f"Attached {Path(display_path).name} — {sent}{scaled}, "
        f"{len(blob) // 1024} KB. The image itself follows in the next message; "
        f"read it there.",
        True,
    )


def _glob(pattern: str, working_dir: str, path: str = None) -> tuple[str, bool]:
    _sess = _remote()
    if _sess is not None:
        return _remote_glob(_sess, pattern, working_dir, path)
    # If pattern is an absolute path, split into root + relative pattern
    p = Path(pattern)
    if p.is_absolute():
        # Find the first glob character to split root from pattern
        parts = p.parts
        root_parts = []
        glob_parts = []
        in_glob = False
        for part in parts:
            if in_glob or any(c in part for c in ("*", "?", "[")):
                in_glob = True
                glob_parts.append(part)
            else:
                root_parts.append(part)
        root = Path(*root_parts) if root_parts else Path("/")
        pattern = str(Path(*glob_parts)) if glob_parts else "**/*"
    else:
        root = _resolve(path, working_dir) if path else Path(working_dir)
    matches = sorted(str(p) for p in root.glob(pattern))
    if not matches:
        return f"No files matching '{pattern}'", True
    # Show relative paths when possible
    result = []
    for m in matches[:200]:
        try:
            result.append(str(Path(m).relative_to(working_dir)))
        except ValueError:
            result.append(m)
    output = "\n".join(result)
    if len(matches) > 200:
        output += f"\n... and {len(matches) - 200} more"
    return output, True


def _grep(pattern: str, working_dir: str, path: str = None, glob: str = None, case_insensitive: bool = False) -> tuple[str, bool]:
    _sess = _remote()
    if _sess is not None:
        return _remote_grep(_sess, pattern, working_dir, path, glob, case_insensitive)
    cmd = ["grep", "-rn", "--include=" + (glob or "*")]
    if case_insensitive:
        cmd.append("-i")
    cmd.append(pattern)
    cmd.append(path if path else working_dir)

    try:
        result = subprocess.run(cmd, capture_output=True, text=True, cwd=working_dir, timeout=30)
        output = result.stdout or result.stderr or "No matches found"
        if len(output) > 8000:
            lines = output.splitlines()
            output = "\n".join(lines[:150]) + f"\n... ({len(lines)} total matches, showing first 150)"
        return output, True
    except subprocess.TimeoutExpired:
        return "grep timed out", False


def _list_dir(working_dir: str, path: str = None) -> tuple[str, bool]:
    _sess = _remote()
    if _sess is not None:
        return _remote_list_dir(_sess, working_dir, path)
    target = _resolve(path, working_dir) if path else Path(working_dir)
    if not target.exists():
        return f"Directory not found: {path}", False
    if not target.is_dir():
        return f"Not a directory: {path}", False

    entries = sorted(target.iterdir(), key=lambda p: (p.is_file(), p.name.lower()))
    lines = []
    for entry in entries:
        if entry.is_dir():
            lines.append(f"  {entry.name}/")
        else:
            size = entry.stat().st_size
            if size < 1024:
                sz = f"{size}B"
            elif size < 1024 * 1024:
                sz = f"{size // 1024}KB"
            else:
                sz = f"{size // (1024*1024)}MB"
            lines.append(f"  {entry.name:<40} {sz:>8}")

    header = f"{target}/\n"
    return header + "\n".join(lines), True


def _web_search(query: str, max_results: int = 10, region: str = "wt-wt") -> tuple[str, bool]:
    if not _HAS_DDG:
        return "duckduckgo-search package not installed. Run: pip install duckduckgo-search", False

    max_results = min(max_results, 20)
    try:
        with _DDGS(timeout=20) as ddgs:
            results = list(ddgs.text(query, region=region, max_results=max_results))
    except Exception as e:
        return f"Search failed: {e}", False

    if not results:
        return "No results found.", True

    lines = [f"Search results for: {query}\n"]
    for i, r in enumerate(results, 1):
        title = r.get("title", "(no title)")
        url = r.get("href", "")
        body = r.get("body", "")
        lines.append(f"[{i}] {title}")
        lines.append(f"    URL: {url}")
        if body:
            snippet = body[:200] + ("…" if len(body) > 200 else "")
            lines.append(f"    {snippet}")
        lines.append("")

    return "\n".join(lines), True


_MAX_DOWNLOAD_BYTES = 1 * 1024 * 1024 * 1024  # 1 GB hard limit


def _web_fetch(url: str, max_chars: int = 8000) -> tuple[str, bool]:
    if not _HAS_REQUESTS:
        return "requests package not installed. Run: pip install requests", False

    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/120.0.0.0 Safari/537.36"
        )
    }
    try:
        # HEAD request first — check Content-Length before downloading anything
        try:
            head = _requests.head(url, headers=headers, timeout=10,
                                  allow_redirects=True)
            content_length = int(head.headers.get("content-length", 0))
            if content_length > _MAX_DOWNLOAD_BYTES:
                size_gb = content_length / (1024 ** 3)
                return (
                    f"Refused to download {url} — "
                    f"Content-Length is {size_gb:.2f} GB, exceeds the 1 GB limit.",
                    False,
                )
        except Exception:
            pass  # HEAD not supported by all servers — proceed with GET

        resp = _requests.get(url, headers=headers, timeout=30,
                             allow_redirects=True, stream=True)
        resp.raise_for_status()

        # Check Content-Length from GET response headers too
        content_length = int(resp.headers.get("content-length", 0))
        if content_length > _MAX_DOWNLOAD_BYTES:
            resp.close()
            size_gb = content_length / (1024 ** 3)
            return (
                f"Refused to download {url} — "
                f"Content-Length is {size_gb:.2f} GB, exceeds the 1 GB limit.",
                False,
            )

        # Read with a rolling size guard (handles chunked/no Content-Length)
        chunks = []
        downloaded = 0
        for chunk in resp.iter_content(chunk_size=1024 * 1024):  # 1 MB chunks
            downloaded += len(chunk)
            if downloaded > _MAX_DOWNLOAD_BYTES:
                resp.close()
                return (
                    f"Aborted download of {url} — "
                    f"exceeded 1 GB limit after {downloaded / (1024**3):.2f} GB.",
                    False,
                )
            chunks.append(chunk)
        resp._content = b"".join(chunks)
    except _requests.exceptions.HTTPError as e:
        status = getattr(e.response, "status_code", None)
        if status == 403:
            return (
                f"Fetch failed: 403 Forbidden — {url}\n"
                "The page requires login or is behind a paywall. "
                "Try: (1) searching for a preprint or open-access version, "
                "(2) fetching https://web.archive.org/web/*/" + url + " for a cached copy, "
                "or (3) using web_search to find a summary of the content.",
                False,
            )
        return f"Fetch failed: {e}", False
    except Exception as e:
        return f"Fetch failed: {e}", False

    content_type = resp.headers.get("content-type", "")

    # PDF — extract text with pymupdf if available, else warn
    if "application/pdf" in content_type or url.lower().endswith(".pdf"):
        try:
            import fitz  # pymupdf
            doc = fitz.open(stream=resp.content, filetype="pdf")
            pages = []
            for i, page in enumerate(doc):
                pages.append(f"--- Page {i+1} ---\n{page.get_text()}")
            text = "\n\n".join(pages)
            if len(text) > max_chars:
                text = text[:max_chars] + f"\n\n… [truncated, {len(text)} total chars]"
            return f"PDF: {url}\n\n{text}", True
        except ImportError:
            return (
                f"Skipped {url} — PDF detected but pymupdf is not installed. "
                "Run: pip install pymupdf",
                False,
            )
        except Exception as e:
            return f"Failed to extract PDF text from {url}: {e}", False

    # Reject other binary content types
    _BINARY_TYPES = ("application/octet-stream", "image/", "audio/",
                     "video/", "application/zip", "application/x-", "application/vnd.")
    if any(content_type.startswith(bt) for bt in _BINARY_TYPES):
        return (
            f"Skipped {url} — binary content ({content_type}), cannot read.",
            False,
        )

    # Plain text / markdown / code — return as-is
    if "text/plain" in content_type or url.endswith((".md", ".txt", ".rst", ".py", ".js", ".ts")):
        text = resp.text[:max_chars]
        if len(resp.text) > max_chars:
            text += f"\n\n… [truncated, {len(resp.text)} total chars]"
        return text, True

    # HTML — extract main text via BeautifulSoup
    if _HAS_BS4:
        soup = _BS4(resp.text, "html.parser")
        # Remove noise elements
        for tag in soup(["script", "style", "nav", "footer", "header", "aside",
                          "form", "noscript", "iframe", "svg", "button"]):
            tag.decompose()
        # Try to find main content block
        main = (
            soup.find("main")
            or soup.find("article")
            or soup.find(id="content")
            or soup.find(class_="content")
            or soup.find("body")
            or soup
        )
        text = main.get_text(separator="\n", strip=True)
        # Collapse excessive blank lines
        import re
        text = re.sub(r"\n{3,}", "\n\n", text)
    else:
        # Fallback: crude tag stripping
        import re
        text = re.sub(r"<[^>]+>", " ", resp.text)
        text = re.sub(r"\s+", " ", text).strip()

    if len(text) > max_chars:
        text = text[:max_chars] + f"\n\n… [truncated, {len(text)} total chars]"

    return f"URL: {url}\n\n{text}", True


def _ensure_playwright() -> bool:
    """Install playwright + chromium if not present. Returns True if available after attempt."""
    global _HAS_PLAYWRIGHT, _sync_playwright  # noqa: PLW0603
    if _HAS_PLAYWRIGHT:
        return True
    try:
        import sys
        print("[crawl_tree] Playwright not found — installing automatically…")
        subprocess.run(
            [sys.executable, "-m", "pip", "install", "playwright", "-q"],
            check=True, capture_output=True,
        )
        subprocess.run(
            [sys.executable, "-m", "playwright", "install", "chromium", "--with-deps"],
            check=True, capture_output=True,
        )
        from playwright.sync_api import sync_playwright as _sp
        _sync_playwright = _sp  # type: ignore[assignment]
        _HAS_PLAYWRIGHT = True
        print("[crawl_tree] Playwright installed successfully.")
        return True
    except Exception as e:
        print(f"[crawl_tree] Auto-install failed: {e} — falling back to requests+BS4.")
        return False


def _crawl_tree(
    root_url: str,
    link_selector: str = "a",
    url_pattern: str = "",
    max_depth: int = 3,
    max_pages: int = 100,
    same_domain: bool = True,
    use_js: bool = False,
    output_path: str = "",
) -> tuple[str, bool]:
    """BFS tree crawl. Returns JSON: {url: {title, text_snippet, children, depth}}"""
    import json
    import re
    from collections import deque
    from urllib.parse import urljoin, urlparse

    url_re = re.compile(url_pattern) if url_pattern else None
    origin = urlparse(root_url).netloc
    visited: dict[str, dict] = {}
    queue: deque[tuple[str, int]] = deque([(root_url, 0)])

    _HEADERS = {
        "User-Agent": (
            "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/120.0.0.0 Safari/537.36"
        )
    }

    # Auto-install Playwright on first use; fall back to requests if it fails
    playwright_ok = _ensure_playwright()
    use_playwright = playwright_ok and (use_js or not _HAS_REQUESTS)

    if use_playwright:
        browser_ctx = _sync_playwright().__enter__()
        browser = browser_ctx.chromium.launch(headless=True)
        page = browser.new_page()
    else:
        browser_ctx = browser = page = None  # type: ignore

    def _fetch_html(url: str) -> str | None:
        try:
            if use_playwright and page is not None:
                page.goto(url, timeout=15000, wait_until="domcontentloaded")
                try:
                    page.wait_for_load_state("networkidle", timeout=5000)
                except Exception:
                    pass
                return page.content()
            else:
                if not _HAS_REQUESTS:
                    return None
                r = _requests.get(url, headers=_HEADERS, timeout=15, allow_redirects=True)
                r.raise_for_status()
                return r.text
        except Exception:
            return None

    def _extract(html: str, base_url: str) -> tuple[str, str, list[str]]:
        """Returns (title, text_snippet, child_urls)."""
        if not _HAS_BS4:
            import re as _re
            title = (_re.search(r"<title[^>]*>(.*?)</title>", html, _re.I | _re.S) or ["", ""])[1]
            text = _re.sub(r"<[^>]+>", " ", html)
            text = " ".join(text.split())[:300]
            hrefs = _re.findall(r'href=["\']([^"\']+)["\']', html)
            return str(title), text, [urljoin(base_url, h) for h in hrefs]

        soup = _BS4(html, "html.parser")
        title = soup.title.get_text(strip=True) if soup.title else base_url
        for tag in soup(["script", "style", "nav", "footer", "header", "aside", "noscript"]):
            tag.decompose()
        body = soup.find("main") or soup.find("article") or soup.find("body") or soup
        snippet = " ".join(body.get_text(separator=" ", strip=True).split())[:300]
        links = []
        for a in soup.select(link_selector):
            href = a.get("href", "")
            if not href or href.startswith(("#", "mailto:", "javascript:", "tel:")):
                continue
            links.append(urljoin(base_url, href).split("#")[0])
        return title, snippet, links

    try:
        while queue and len(visited) < max_pages:
            url, depth = queue.popleft()
            if url in visited:
                continue

            html = _fetch_html(url)
            if html is None:
                visited[url] = {"title": "FETCH_ERROR", "text_snippet": "", "children": [], "depth": depth}
                continue

            title, snippet, raw_links = _extract(html, url)

            children: list[str] = []
            if depth < max_depth:
                seen_this_page: set[str] = set()
                for child in raw_links:
                    if child in visited or child in seen_this_page:
                        continue
                    if same_domain and urlparse(child).netloc != origin:
                        continue
                    if url_re and not url_re.search(child):
                        continue
                    seen_this_page.add(child)
                    children.append(child)
                    queue.append((child, depth + 1))

            visited[url] = {
                "title": title,
                "text_snippet": snippet,
                "children": children,
                "depth": depth,
            }
    finally:
        if use_playwright and browser_ctx is not None:
            try:
                browser.close()  # type: ignore
                browser_ctx.__exit__(None, None, None)
            except Exception:
                pass

    result_json = json.dumps(visited, indent=2, ensure_ascii=False)

    if output_path:
        try:
            Path(output_path).parent.mkdir(parents=True, exist_ok=True)
            Path(output_path).write_text(result_json, encoding="utf-8")
            saved_msg = f"\n\nSaved to: {output_path}"
        except Exception as e:
            saved_msg = f"\n\nFailed to save: {e}"
    else:
        saved_msg = ""

    engine = "Playwright" if use_playwright else "requests+BS4"
    summary = (
        f"Crawled {len(visited)} pages from {root_url} "
        f"(depth≤{max_depth}, engine={engine}){saved_msg}\n\n"
        + result_json
    )
    return summary, True
